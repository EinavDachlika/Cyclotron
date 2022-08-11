from tkinter import *
from tkinter import ttk
from PIL import Image, ImageTk
import mysql.connector
from mysql.connector import Error
import pandas as pd
from docx.api import Document
import aspose.words as aw
from tkinter import filedialog as fd

##table code
# https://pythonguides.com/python-tkinter-table-tutorial/

root = Tk()
# root.geometry("300x300")


root.title("Sheri Orders System")#Setting->to Main(Sheri Orders system)

# defult font
root.option_add("*Font", "Helvetica")

# connect to MySqL
try:
    #Maor local DB Mysql
    # db = mysql.connector.connect(
    #     host="localhost",
    #     port=3308,
    #     user="root",
    #     password="root",
    #     database="cyclotron")

    # # Einav local DB-Mysql
    db = mysql.connector.connect(
      host="localhost",
      user="root",
      password="Cyclotron2022@?%",
      database= "cyclotron")

    if db.is_connected():
        # db_Info = db.get_server_info()
        # print("Connected to MySQL Server version ", db_Info)
        dbCursor = db.cursor(buffered=True)
        # Check to see if connection to Mysql was created
        print("connection to local mysql succeed", db)

# dbCursor.execute("select database();")
# record = dbCursor.fetchone()
# print("You're connected to database: ", record)
except Error as e:
    print("Error while connecting to MySQL", e)

######################Hospital page##########################################
hospitalFrame = Frame(root)
h = Scrollbar(hospitalFrame, orient='horizontal')
#hospitalFrame.pack(fill=X)


# feed label
feedLabel = Label(hospitalFrame, text = 'Hospitals Details', font=('Helvetica',26, 'bold'),fg='#034672')
PlaceLable_X=50
PlaceLable_Y=10

feedLabel.pack(side=LEFT)
feedLabel.place(x=PlaceLable_X,y=PlaceLable_Y)


# scrollbar
Cyclotron_scroll = Scrollbar(hospitalFrame ,orient="vertical",width=25)
# Cyclotron_scroll.pack(side=LEFT)
# Cyclotron_scroll.place(x=550, y= 160)

hospitals_list = ttk.Treeview(hospitalFrame, yscrollcommand=Cyclotron_scroll.set,height=12)

hospitals_list.pack(side=LEFT, padx=PlaceLable_X+50, pady=PlaceLable_Y+80)

# column define

hospitals_list['columns'] = ('Name', 'Fixed Activity Level (mci)', 'Transport Time (minutes)')

# column format
width_Version=110
width_Capacity=110
width_Efficiency=185
width_Description=110

hospitals_list.column("#0", width=0, stretch=NO)
hospitals_list.column("Name", anchor=CENTER, width=width_Version)
hospitals_list.column("Fixed Activity Level (mci)", anchor=CENTER, width=width_Capacity)
hospitals_list.column("Transport Time (minutes)", anchor=CENTER, width=width_Efficiency)

# Create Headings
hospitals_list.heading("#0", text="", anchor=CENTER)
hospitals_list.heading("Name", text="Name", anchor=CENTER)
hospitals_list.heading("Fixed Activity Level (mci)", text="Fixed Activity Level (mci)", anchor=CENTER)
hospitals_list.heading("Transport Time (minutes)", text="Transport Time (minutes)", anchor=CENTER)

# add data from db
cursor = db.cursor();
cursor.execute("SELECT * FROM hospital");
hospitals_in_db = cursor.fetchall();

#Insert data of Hospitals into My-SQl DB
#The INSERT IGNORE statement will cause MySQL to do nothing when the insertion throws an error. If there’s no error, then a new row will be added to the table.
#cursor.execute("INSERT IGNORE INTO hospital (idhospital,Name,Fixed_activity_level,Transport_time) VALUES (1,'Belinson',9.2,15.0),(2,'Ichilov',10.0,20.0),(3,'Assuta TA',10.9,30.0),(4,'Sheb',10.5,35.0),(5,'Ziv',11.0,25.0),(6,'Assuta Ashdod',13.1,60.0),(7,'Assaf Harofeh',10.6,65.0),(8,'Augusta Victoria',9.6,50.0),(9,'Hila Pharma',9.6,50.0),(10,'Hadassah',9.5,0.0);")

#Insert 2 material to the DB,material table
#cursor.execute("INSERT IGNORE INTO material (idmaterial,materialName) VALUES (1,'FDG'),(2,'FDOPA');")
#cleanup DB
db.commit();
# cursor.close();
# db.close();

iid=0
for hospital in hospitals_in_db:
    #print(hospital)
    hospitals_list.insert(parent='', index='end', iid=iid, text='',
                          values=(hospital[1], hospital[2], hospital[3]))
    iid +=1

hospitals_list.pack()



# def open_popup_hospital():
#     pass
#
# def delete_hospital():
#     pass

#Create a button in the main Window to open the popup
editIcon = Image.open("editIcon.jpg")
resizedEditIcon = editIcon.resize((20,20), Image.ANTIALIAS)
imgEdit = ImageTk.PhotoImage(resizedEditIcon)
editButton=Button(hospitalFrame, image=imgEdit, borderwidth=0)
editButton.pack()
editButton.place(x=425, y=55)


# delete button (Icon) - List
deleteIcon = Image.open("‏‏deleteIcon.png")
resizedDeleteIcon = deleteIcon.resize((20,20), Image.ANTIALIAS)
imgDelete = ImageTk.PhotoImage(resizedDeleteIcon)
deleteButton=Button(hospitalFrame, image=imgDelete, borderwidth=0)
deleteButton.pack()
deleteButton.place(x=470, y=55)

#####################end of hospital page###################################################

def hospital_page():
    """ this function is swap function for hospital frame/page"""
    settingButton.config(bg='#F0F0F0');
    ordersButton.config(bg='#F0F0F0');
    hospitalsButton.config(bg='gray');
    hospitalFrame.pack(fill='both',expand=1);
    SettingsFrame.forget();
    ordersFrame.forget();

##################Orders page############################################

ordersFrame = Frame(root)
#h = Scrollbar(ordersFrame, orient='horizontal')
#ordersFrame.pack(fill=X)


# feed label
feedLabel = Label(ordersFrame, text ='Orders', font=('Helvetica', 26, 'bold'), fg='#034672')
PlaceLable_X=50
PlaceLable_Y=10

feedLabel.pack(side=LEFT)
feedLabel.place(x=PlaceLable_X,y=PlaceLable_Y)


# scrollbar
Cyclotron_scroll = Scrollbar(ordersFrame, orient="vertical", width=25)
# Cyclotron_scroll.pack(side=LEFT)
# Cyclotron_scroll.place(x=550, y= 160)

OrdersTree = ttk.Treeview(ordersFrame, yscrollcommand=Cyclotron_scroll.set, height=12)

OrdersTree.pack(side=LEFT, padx=PlaceLable_X+50, pady=PlaceLable_Y+80)
my_label=Label(root,text='');

def openFile():
    """This is function for importing order files"""
    filename = fd.askopenfilename(
        initialdir="D:\PythonProjects\Cyclotron",
        title="Open a file",
        filetype=(("Word files","*.docx"),("Word files","*.doc"),("xlsx files","*.xlsx"),("All Files","*.*"),("PDF files","*.pdf"))
    )


    if filename :
        if  "xlsx" in filename :#Excel file
            try:
                filename=r"{}".format(filename)
                df=pd.read_excel(filename)
            except ValueError:
                my_label.config(text="File couldn't be open,try again");
            except FileNotFoundError:
                my_label.config(text="File couldn't be open,try again");

            clear_tree();

            OrdersTree["column"] =  list(df.columns);
            OrdersTree["show"] = "headings";

            for column in OrdersTree["column"]:
                OrdersTree.heading(column,text=column)

            df_rows=df.to_numpy().tolist();

            for row in df_rows:
                OrdersTree.insert("","end",values=row)

                OrdersTree.pack();



        if "docx" in filename or "doc" in filename:#word files
            #convert word to excel

            if (("doc" in filename) and ("docx" not in filename)):#convert docx to doc
                doc = aw.Document(filename)
                filename="NewWordOutput1.docx";
                doc.save(filename)


            document = Document(filename)
            tables = document.tables
            df = pd.DataFrame()

            for table in document.tables:
                for row in table.rows:
                    text = [cell.text for cell in row.cells]
                    df = df.append([text], ignore_index=True)

            #df.columns = ["Column1", "Column2","Column3","Column4","Column5","Column6","Column7","Column8"]
            df.to_excel("D:/PythonProjects/Cyclotron/OrderOutputTest.xlsx")
            #print(df);


            clear_tree();

            OrdersTree["column"] =  list(df.columns);
            OrdersTree["show"] = "headings";

            for column in OrdersTree["column"]:
                OrdersTree.heading(column,text=column)

            df_rows=df.to_numpy().tolist();

            for row in df_rows:
                OrdersTree.insert("","end",values=row)

                OrdersTree.pack();




def clear_tree():
    OrdersTree.delete(*OrdersTree.get_children())





#Create a button in the main Window to open the popup
editIcon = Image.open("editIcon.jpg")
resizedEditIcon = editIcon.resize((20,20), Image.ANTIALIAS)
imgEdit = ImageTk.PhotoImage(resizedEditIcon)
editButton=Button(ordersFrame, image=imgEdit, borderwidth=0)
editButton.pack()
editButton.place(x=425, y=55)

#Create a button for import order from files(Excel or Word)
ImportFileIcon = Image.open("ImportFile2.png")
resized_Edit_Icon = ImportFileIcon.resize((80,20), Image.ANTIALIAS)
img_Edit = ImageTk.PhotoImage(resized_Edit_Icon)
importFileButton=Button(ordersFrame, image=img_Edit, borderwidth=0,command=openFile)
importFileButton.pack()
importFileButton.place(x=230, y=65)


# edit_button = Button(hospitalFrame, text= "Edit", command= open_popup_hospital)
# edit_button.pack(side= LEFT)
# edit_button.place(x=450, y=50)


# delete button (Icon) - List
deleteIcon = Image.open("‏‏deleteIcon.png")
resizedDeleteIcon = deleteIcon.resize((20,20), Image.ANTIALIAS)
imgDelete = ImageTk.PhotoImage(resizedDeleteIcon)
deleteButton=Button(ordersFrame, image=imgDelete, borderwidth=0)
deleteButton.pack()
deleteButton.place(x=470, y=55)

##############end of Orders page#########################################

def Orders_page():
    """ this function is swap function for Orders frame/page"""
    settingButton.config(bg='#F0F0F0');  ##F0F0F0 is default color(gray)
    hospitalsButton.config(bg='#F0F0F0');
    ordersButton.config(bg="gray");
    ordersFrame.pack(fill='both',expand=1);
    hospitalFrame.forget();
    SettingsFrame.forget();




##################### toolbar #####################
toolbarbgcolor = "white"
toolbar = Frame(root, bg=toolbarbgcolor)
#toolbar.grid(sticky='nesw')

# add logo - toolbar
LogoImagePath = Image.open("LogoImage.png")
LogoImageResize = LogoImagePath.resize((120, 57), Image.ANTIALIAS)
LogoImage = ImageTk.PhotoImage(LogoImageResize)
Label(toolbar, image=LogoImage).pack(side=LEFT, padx=10, pady=6)

# work plan button - toolbar
workPlanButton = Button(toolbar, text="Work Plans", font='Helvetica 11')
workPlanButton.pack(side=LEFT, padx=10, pady=3)

# Hospitals button - toolbar
hospitalsButton = Button(toolbar, text="Hospitals",command=hospital_page,font='Helvetica 11', activebackground='red')
hospitalsButton.pack(side=LEFT, padx=10, pady=3)

# Orders button - toolbar
ordersButton = Button(toolbar, text="Orders", font='Helvetica 11',command=Orders_page)
ordersButton.pack(side=LEFT, padx=10, pady=3)

# Reports button - toolbar
reportsButton = Button(toolbar, text="Reports", font='Helvetica 11')
reportsButton.pack(side=LEFT, padx=10, pady=3)



#########################################################################
def setting_page():
    """ this function is swap function for viewing setting frame/page"""
    hospitalsButton.config(bg='#F0F0F0')
    ordersButton.config(bg='#F0F0F0');
    settingButton.config(bg="gray");
    SettingsFrame.pack(fill='both',expand=1);
    hospitalFrame.forget();
    ordersFrame.forget();

# settings Icon - toolbar

settingsIcon = Image.open("gearIcon.png")
resizedSettingsIcon = settingsIcon.resize((35, 35), Image.ANTIALIAS)
imgSettings = ImageTk.PhotoImage(resizedSettingsIcon)
settingButton=Button(toolbar,command=setting_page ,image=imgSettings,activebackground='red', borderwidth=0)
settingButton.pack(side=RIGHT, padx=10, pady=3)


toolbar.pack(side=TOP, fill=X)

toolbar.grid_columnconfigure(1, weight=1)
#####################################Setting frame####################


SettingsFrame = Frame(root)
#h = Scrollbar(SettingsFrame, orient='horizontal')
SettingsFrame.pack(fill='both',expand=1)

# feed label
feedLabel = Label(SettingsFrame, text='Settings', font=('Helvetica', 26, 'bold'), fg='#034672')
PlaceLable_X = 50
PlaceLable_Y = 10

feedLabel.pack(side=LEFT)
feedLabel.place(x=PlaceLable_X, y=PlaceLable_Y)

##################### Cyclotron #####################
# Cyclotron Details label
CyclotronLabel = Label(SettingsFrame, text='Cyclotron Details', font=('Helvetica', 15, 'bold'), fg='#034672')
cyclo_Lable_place_x = 80
cyclo_Lable_place_y = 70

CyclotronLabel.pack(side=LEFT)
CyclotronLabel.place(x=cyclo_Lable_place_x, y=cyclo_Lable_place_y)

# scrollbar
Cyclotron_scroll = Scrollbar(SettingsFrame, orient="vertical", width=20)
Cyclotron_scroll.pack(side=LEFT)
Cyclotron_scroll.place(x=613, y=160)

cyclo_list = ttk.Treeview(SettingsFrame, yscrollcommand=Cyclotron_scroll.set, height=5)

cyclo_list.pack(side=LEFT, padx=cyclo_Lable_place_x + 30, pady=cyclo_Lable_place_y + 50)

# Cyclotron_scroll.config(command=cyclo_list.yview)
# Cyclotron_scroll.config(command=cyclo_list.xview)

# column define

cyclo_list['columns'] = ('Version', 'Capacity (mci/h)', 'Constant Efficiency (mCi/mA)', 'Description')

# column format
width_Version = 90
width_Capacity = 110
width_Efficiency = 185
width_Description = 110

cyclo_list.column("#0", width=0, stretch=NO)
cyclo_list.column("Version", anchor=CENTER, width=width_Version)
cyclo_list.column("Capacity (mci/h)", anchor=CENTER, width=width_Capacity)
cyclo_list.column("Constant Efficiency (mCi/mA)", anchor=CENTER, width=width_Efficiency)
cyclo_list.column("Description", anchor=CENTER, width=width_Description)

# Create Headings
cyclo_list.heading("#0", text="", anchor=CENTER)
cyclo_list.heading("Version", text="Version", anchor=CENTER)
cyclo_list.heading("Capacity (mci/h)", text="Capacity (mci/h)", anchor=CENTER)
cyclo_list.heading("Constant Efficiency (mCi/mA)", text="Constant Efficiency (mCi/mA)", anchor=CENTER)
cyclo_list.heading("Description", text="Description", anchor=CENTER)

# add data from db
cursor = db.cursor()
# Test to see if DB was created
# cursor.execute("SHOW DATABASES")
# for cyclotron in cursor:
#     print(cyclotron)

cursor.execute("SELECT * FROM resourcecyclotron");
cyclotrons = cursor.fetchall();

#Insert data of Settings page into My-SQl
#The INSERT IGNORE statement will cause MySQL to do nothing when the insertion throws an error. If there’s no error, then a new row will be added to the table.
cursor.execute("INSERT IGNORE INTO resourcecyclotron (idresourceCyclotron,version,capacity,constant_efficiency,description) VALUES (1,2.1,2000,220,'site1'),(2,2.2,1700,150,'site2');")
#cleanup
db.commit()
# cursor.close()
# db.close()


iid = 0
for cyclo in cyclotrons:
    print(cyclo)
    cyclo_list.insert(parent='', index='end', iid=iid, text='',
                      values=(cyclo[1], cyclo[2], cyclo[3], cyclo[4]))
    iid += 1

cyclo_list.pack()

# frame = Frame(root)
# frame.pack()

get_version = ""
get_capacity = ""
get_efficiency = ""
get_description = ""


class Cyclotron:
    def _init_(self, version, capacity, constant_efficiency, description):
        self.version = version
        self.capacity = capacity
        self.constant_efficiency = constant_efficiency
        self.description = description

    # def edit(self):
    #    def update_record(get_version, get_capacity, get_efficiency, get_description):
    #       selected = cyclo_list.focus()
    #       # save new data
    #       print("get_version" + get_version)
    #       cyclo_list.item(selected, text="", values=(get_version, get_capacity, get_efficiency, get_description))
    #
    #       # # clear entry boxes
    #       # Version_entry.delete(0, END)
    #       # Capacity_entry.delete(0, END)
    #       # Efficiency_entry.delete(0, END)


def open_popup_cyclotron():
    edit_popup = Toplevel(root)
    edit_popup.geometry("900x400")
    edit_popup.title("Edit Cyclotron Details")
    Label(edit_popup, text="Edit Cyclotron Details", font=('Helvetica 17 bold'), fg='#034672').place(x=10, y=18)

    # labels
    popup_label_y = 80
    Version = Label(edit_popup, text="Version")
    Version.grid(row=1, column=1)
    version_x = 20
    Version.place(x=version_x, y=popup_label_y)

    Capacity = Label(edit_popup, text="Capacity")
    Capacity_units = Label(edit_popup, text="(mci/h)")
    Capacity_units.config(font=("Courier", 9))
    Capacity.grid(row=1, column=2)
    capacity_x = version_x + Version.winfo_reqwidth() + 70
    Capacity.place(x=capacity_x, y=popup_label_y)
    capacity_units_x = capacity_x + Capacity.winfo_reqwidth()
    Capacity_units.place(x=capacity_units_x, y=popup_label_y + 7)

    Efficiency = Label(edit_popup, text="Constant Efficiency")
    Efficiency_units = Label(edit_popup, text="(mCi/mA)")
    Efficiency_units.config(font=("Courier", 9))
    Efficiency.grid(row=1, column=3)
    efficiency_x = capacity_units_x + Capacity_units.winfo_reqwidth() + 50
    Efficiency.place(x=efficiency_x, y=popup_label_y)
    efficiency_units_x = efficiency_x + Efficiency.winfo_reqwidth()
    Efficiency_units.place(x=efficiency_units_x, y=popup_label_y + 7)

    Description = Label(edit_popup, text="Description")
    Description.grid(row=1, column=3)
    description_x = efficiency_units_x + Efficiency_units.winfo_reqwidth() + 30
    Description.place(x=description_x, y=popup_label_y)

    # Entry boxes
    Version_entry = Entry(edit_popup, width=10)
    Version_entry.grid(row=2, column=1)
    Version_entry.place(x=version_x + 3, y=popup_label_y + 30)

    Capacity_entry = Entry(edit_popup, width=14)
    Capacity_entry.grid(row=2, column=2)
    Capacity_entry.place(x=capacity_x, y=popup_label_y + 30)

    Efficiency_entry = Entry(edit_popup, width=15)
    Efficiency_entry.grid(row=2, column=3)
    Efficiency_entry.place(x=efficiency_x, y=popup_label_y + 30)

    Description_entry = Entry(edit_popup, width=15)
    Description_entry.grid(row=2, column=4)
    Description_entry.place(x=description_x, y=popup_label_y + 30)

    # # clear entry boxes
    # Version_entry.delete(0, END)
    # Capacity_entry.delete(0, END)
    # Efficiency_entry.delete(0, END)

    # grab record
    selected = cyclo_list.focus()
    # grab record values
    values = cyclo_list.item(selected, 'values')
    # temp_label.config(text=selected)

    # insert cyclotron details from db to entry boxes
    Version_entry.insert(0, values[0])
    Capacity_entry.insert(0, values[1])
    Efficiency_entry.insert(0, values[2])
    Description_entry.insert(0, values[3])

    # get_version = Version_entry.get()
    # get_capacity = Capacity_entry.get()
    # get_efficiency = Efficiency_entry.get()
    # get_description = Description_entry.get()

    select_button = Button(edit_popup, text="Save Changes",
                           command=lambda: update_record(Version_entry.get(), Capacity_entry.get(),
                                                         Efficiency_entry.get(), Description_entry.get()))
    select_button.pack(side=LEFT)
    select_button.place(x=370, y=250)


# in the class?
def update_record(get_version, get_capacity, get_efficiency, get_description):
    print("get_version" + get_version)
    selected = cyclo_list.focus()
    print(cyclo_list.item(selected, 'values'))
    # save new data
    cyclo_list.item(selected, text="", values=(get_version, get_capacity, get_efficiency, get_description))

    # # clear entry boxes
    # Version_entry.delete(0, END)
    # Capacity_entry.delete(0, END)
    # Efficiency_entry.delete(0, END)


# # save Record
# def update_record():
#     selected = cyclo_list.focus()
#     # save new data
#     cyclo_list.item(selected, text="", values=(Version_entry.get(), Capacity_entry.get(), Efficiency_entry.get()))
#
#     # clear entry boxes
#     Version_entry.delete(0, END)
#     Capacity_entry.delete(0, END)
#     Efficiency_entry.delete(0, END)


# Create a button in the main Window to edit  record (open the popup) - cyclotron
cyclotronEditIcon = Image.open("editIcon.jpg")
resizedCycloEditIcon = cyclotronEditIcon.resize((20, 20), Image.ANTIALIAS)
imgEditCyclotron = ImageTk.PhotoImage(resizedCycloEditIcon)
editCyclotronButton = Button(SettingsFrame, image=imgEditCyclotron, borderwidth=0, command=open_popup_cyclotron)
editCyclotronButton.pack(side=LEFT)
editCyclotronButton.place(x=cyclo_Lable_place_x + 450, y=cyclo_Lable_place_y + 15)

# edit_button = Button(SettingsFrame, text= "Edit", command= open_popup_cyclotron, width=4, height=1)
# edit_button.pack(side= LEFT)
# edit_button.place(x=270, y=265)


# Create a button in the main Window to Delete record - cyclotron
cyclotronDeleteIcon = Image.open("‏‏deleteIcon.png")
resizedCycloDeleteIcon = cyclotronDeleteIcon.resize((20, 20), Image.ANTIALIAS)
imgDeleteCyclotron = ImageTk.PhotoImage(resizedCycloDeleteIcon)
deleteCyclotronButton = Button(SettingsFrame, image=imgDeleteCyclotron, borderwidth=0, command=open_popup_cyclotron)
deleteCyclotronButton.pack(side=LEFT)
deleteCyclotronButton.place(x=cyclo_Lable_place_x + 500, y=cyclo_Lable_place_y + 15)

# Create a button in the main Window to add record - cyclotron
cyclotronAddIcon = Image.open("addIcon.png")
resizedCycloAddIcon = cyclotronAddIcon.resize((25, 25), Image.ANTIALIAS)
imgAddCyclotron = ImageTk.PhotoImage(resizedCycloAddIcon)
addCyclotronButton = Button(SettingsFrame, image=imgAddCyclotron, borderwidth=0, command=open_popup_cyclotron)
addCyclotronButton.pack(side=LEFT)
addCyclotronButton.place(x=cyclo_Lable_place_x + 400, y=cyclo_Lable_place_y + 14)

# add_button = Button(SettingsFrame, text="Add Cyclotron", command= open_popup_cyclotron, width = 4, height=1)
# add_button.pack(side= LEFT)
# add_button.place(x=370, y=265)




##################### Module #####################
# Module Details label
moduleLabel = Label(SettingsFrame, text = 'Module Details', font=('Helvetica',15, 'bold'),fg='#034672')
module_Lable_place_x=700
module_Lable_place_y=70

moduleLabel.pack(side=RIGHT)
moduleLabel.place(x=module_Lable_place_x,y=module_Lable_place_y)


# scrollbar
Module_scroll = Scrollbar(SettingsFrame ,orient="vertical",width=20)
Module_scroll.pack(side=RIGHT)
Module_scroll.place(x=1035, y= 160)

module_list = ttk.Treeview(SettingsFrame, yscrollcommand=Module_scroll.set,height=5)

module_list.pack(side=LEFT, padx=0, pady=module_Lable_place_y+50)



# Module_scroll.config(command=cyclo_list.yview)
# Module_scroll.config(command=cyclo_list.xview)

# column define

module_list['columns'] = ('Version', 'Capacity (mci/h)', 'Description')


module_list.column("#0", width=0, stretch=NO)
module_list.column("Version", anchor=CENTER, width=width_Version)
module_list.column("Capacity (mci/h)", anchor=CENTER, width=width_Capacity)
module_list.column("Description", anchor=CENTER, width=width_Description)

# Create Headings
module_list.heading("#0", text="", anchor=CENTER)
module_list.heading("Version", text="Version", anchor=CENTER)
module_list.heading("Capacity (mci/h)", text="Capacity (mci/h)", anchor=CENTER)
module_list.heading("Description", text="Description", anchor=CENTER)

# add data from db
cursor = db.cursor()
cursor.execute("SELECT * FROM resourcemodule")
modules = cursor.fetchall()

iid=0
for module in modules:
    print(module)
    cyclo_list.insert(parent='', index='end', iid=iid, text='',
               values=(module[1], module[2], module[3]))
    iid +=1

module_list.pack()

get_version=""
get_capacity=""
get_efficiency=""
get_description=""

def open_popup_module():
   edit_popup= Toplevel(root)
   edit_popup.geometry("900x400")
   edit_popup.title("Edit Module Details")
   Label(edit_popup, text= "Edit Module Details", font=('Helvetica 17 bold'), fg='#034672').place(x=10,y=18)

   # labels
   popup_label_y=80
   Version = Label(edit_popup, text="Version")
   Version.grid(row=1, column=1)
   version_x = 20
   Version.place(x=version_x, y=popup_label_y)


   Capacity = Label(edit_popup, text="Capacity")
   Capacity_units = Label(edit_popup, text="(mci/h)")
   Capacity_units.config(font=("Courier", 9))
   Capacity.grid(row=1, column=2)
   capacity_x = version_x+Version.winfo_reqwidth()+70
   Capacity.place(x=capacity_x, y=popup_label_y)
   capacity_units_x=capacity_x + Capacity.winfo_reqwidth()
   Capacity_units.place(x=capacity_units_x, y=popup_label_y+7)


   Description = Label(edit_popup, text="Description")
   Description.grid(row=1, column=3)
   description_x = capacity_units_x + Capacity_units.winfo_reqwidth() + 50
   Description.place(x=description_x, y=popup_label_y)

   # Entry boxes
   Version_entry = Entry(edit_popup, width=10)
   Version_entry.grid(row=2, column=1)
   Version_entry.place(x=version_x+3, y=popup_label_y+30)

   Capacity_entry = Entry(edit_popup, width=14)
   Capacity_entry.grid(row=2, column=2)
   Capacity_entry.place(x=capacity_x, y=popup_label_y+30)


   Description_entry = Entry(edit_popup,width=15)
   Description_entry.grid(row=2, column=4)
   Description_entry.place(x=description_x, y=popup_label_y+30)


   # clear entry boxes
   Version_entry.delete(0, END)
   Capacity_entry.delete(0, END)
   Description_entry.delete(0, END)

   # grab record
   selected = module_list.focus()
   # grab record values
   values = module_list.item(selected, 'values')
   # temp_label.config(text=selected)

   # output to entry boxes
   Version_entry.insert(0, values[0])
   Capacity_entry.insert(0, values[1])
   Description_entry.insert(0, values[2])

   get_version = Version_entry.get()
   print(get_version)
   get_capacity = Capacity_entry.get()
   get_description = Description_entry.get()

   select_button = Button(edit_popup, text="Save Changes", command=update_record)
   select_button.pack(side=LEFT)
   select_button.place(x=370, y=250)


#Create a button in the main Window to edit  record (open the popup) - module
moduleEditIcon = Image.open("editIcon.jpg")
resizedModuleEditIcon = moduleEditIcon.resize((20, 20), Image.ANTIALIAS)
imgEditModule = ImageTk.PhotoImage(resizedModuleEditIcon)
editModuleButton = Button(SettingsFrame, image=imgEditModule, borderwidth=0, command=open_popup_module)
editModuleButton.pack(side= LEFT)
editModuleButton.place(x=module_Lable_place_x+250, y=module_Lable_place_y+15)

# edit_button = Button(SettingsFrame, text= "Edit", command= open_popup_module)
# edit_button.pack(side= LEFT)
# edit_button.place(x=790, y=270)

#Create a button in the main Window to Delete record - module
moduleDeleteIcon = Image.open("‏‏deleteIcon.png")
resizedModuleDeleteIcon = moduleDeleteIcon.resize((20, 20), Image.ANTIALIAS)
imgDeleteModule = ImageTk.PhotoImage(resizedModuleDeleteIcon)
deleteModuleButton = Button(SettingsFrame, image=imgDeleteModule, borderwidth=0, command=open_popup_module)
deleteModuleButton.pack(side= LEFT)
deleteModuleButton.place(x=module_Lable_place_x+300, y=module_Lable_place_y+15)

#Create a button in the main Window to add record - module
moduleAddIcon = Image.open("addIcon.png")
resizedModuleAddIcon = moduleAddIcon.resize((25, 25), Image.ANTIALIAS)
imgAddModule = ImageTk.PhotoImage(resizedModuleAddIcon)
addModuleButton = Button(SettingsFrame, image=imgAddModule, borderwidth=0, command=open_popup_cyclotron)
addModuleButton.pack(side= LEFT)
addModuleButton.place(x=module_Lable_place_x+200, y=module_Lable_place_y+14)

# add_button = Button(SettingsFrame, text="Add", command= open_popup_cyclotron)
# add_button.pack(side= LEFT)
# add_button.place(x=890, y=270)

def update_record():
    selected = cyclo_list.focus()
    # save new data
    print("get_version"+get_version)
    cyclo_list.item(selected, text="", values=(get_version, get_capacity, get_efficiency, get_description))

    # # clear entry boxes
    # Version_entry.delete(0, END)
    # Capacity_entry.delete(0, END)
    # Efficiency_entry.delete(0, END)


# # # save Record
# # def update_record():
# #     selected = cyclo_list.focus()
# #     # save new data
# #     cyclo_list.item(selected, text="", values=(Version_entry.get(), Capacity_entry.get(), Efficiency_entry.get()))
# #
# #     # clear entry boxes
# #     Version_entry.delete(0, END)
# #     Capacity_entry.delete(0, END)
# #     Efficiency_entry.delete(0, END)


SettingsFrame.pack(fill='both',expand=1)
root.mainloop()
