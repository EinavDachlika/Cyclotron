from tkinter import *
from tkinter import ttk,messagebox
import tkinter as tk
from PIL import Image, ImageTk
import mysql.connector
from mysql.connector import Error
from tkinter import filedialog as fd
import pandas as pd
from docx.api import Document
import aspose.words as aw
import xlrd #Version 1.2.0
from tkcalendar import DateEntry,Calendar
from datetime import date,timedelta,datetime
import csv
import traceback
import logging
import sys,importlib
from ConnectToDB import *          #connect to mysql DB
#import DB_tables                   #create tables ans import them
#import Permission

#import DB_tables_test
# def UsersFunction():
spec = importlib.util.spec_from_file_location("module.name", "D:\PythonProjects\Cyclotron\‏‏DB_tables_test.py")
foo = importlib.util.module_from_spec(spec)
sys.modules["module.name"] = foo
spec.loader.exec_module(foo)

###############################################################################################
root = tk.Tk();
#root.geometry("300x300")


#root.title("Orders")

#defult font
root.option_add("*Font", "Helvetica")

# ##################### toolbar #####################


#############################################
#Disable foreiken-key relations
# cursor = db.cursor();
# Hack_Disabled_ChildSqlForeignKeys="SET FOREIGN_KEY_CHECKS=1;"
#
#
# cursor.execute(Hack_Disabled_ChildSqlForeignKeys);
# db.commit();
# cursor.close();

################################################################
toolbarbgcolor = "white"
toolbar = Frame(root, bg=toolbarbgcolor)
toolbar.grid(sticky='nesw')

# add logo - toolbar
LogoImagePath = Image.open("LogoImage.png")
LogoImageResize = LogoImagePath.resize((120, 57),Image.ANTIALIAS)
LogoImage = ImageTk.PhotoImage(LogoImageResize)
Label(toolbar,image=LogoImage).pack(side=LEFT,padx=10,pady=6)

# work plan button - toolbar
workPlanButton = Button(toolbar, text="Work Plans",font='Helvetica 11')
workPlanButton.pack(side=LEFT,padx=10,pady=3)


# Hospitals button - toolbar
hospitalsButton = Button (toolbar, text="Hospitals",font='Helvetica 11', activebackground='gray')
hospitalsButton.pack(side=LEFT,padx=10,pady=3)

# Orders button - toolbar
ordersButton = Button (toolbar, text="Orders", font='Helvetica 11')
ordersButton.pack(side=LEFT,padx=10,pady=3)


# Reports button - toolbar
reportsButton = Button (toolbar, text="Reports", font='Helvetica 11')
reportsButton.pack(side=LEFT,padx=10,pady=3)

# settings Icon - toolbar

settingsIcon = Image.open("gearIcon.png")
resizedSettingsIcon = settingsIcon.resize((35,35), Image.ANTIALIAS)
imgSettings = ImageTk.PhotoImage(resizedSettingsIcon)
Button(toolbar, image=imgSettings, borderwidth=0).pack(side=RIGHT,padx=10,pady=3)

# # Users Icon - toolbar
def usersFunc():
      pass
#     Permission.OutputScreen();


# LogInIcon = Image.open("./Images/UsersIcon.jpg")
# resizedLogInIcon = LogInIcon.resize((35,35), Image.ANTIALIAS)
# imgLogIn = ImageTk.PhotoImage(resizedLogInIcon)
# Button(toolbar, image=imgLogIn, borderwidth=0,command=usersFunc).pack(side=RIGHT,padx=10,pady=3)


toolbar.pack(side=TOP, fill=X)

toolbar.grid_columnconfigure(1, weight=1)

##################### Start Order page ###################################################

ordersFrame = Frame(root);
ordersFrame.config(bg="#F0F3F4");#color of page-white-gray

ordersFrame.pack(fill=X)

# ScrollbarOrderMainPage
feedLabel = Label(ordersFrame, text ='Orders', font=('Helvetica', 26, 'bold'), fg='#034672');
feedLabel.pack(side=LEFT);
feedLabel.place(x=50,y=10);

# # admin label
# AdminLabelOrderslpage1 = Label(ordersFrame, text=f"{Permission.ValidateTypeOfUser} Connected:", font=('Helvetica', 13, 'bold'), fg='red')
# AdminLabelOrderslpage1.pack();
# AdminLabelOrderslpage1.place(x=350, y=20);
#
# # admin connected label
# NameOfAdminLabeLoRDERSpage2 = Label(ordersFrame, text=Permission.user_verified, font=('Helvetica', 13, 'bold'), fg='red')
# NameOfAdminLabeLoRDERSpage2.pack();
# NameOfAdminLabeLoRDERSpage2.place(x=510, y=20);

# # user label
# UserLabelHospitalpage1 = Label(hospitalFrame, text=f"{Permission.ValidateTypeOfUser} connected:", font=('Helvetica', 13, 'bold'), fg='red')
# UserLabelHospitalpage1.pack();
# UserLabelHospitalpage1.place(x=350, y=20);
#
# # user  connected label
# NameOfUserLabelHospitalpage2 = Label(hospitalFrame, text=Permission.user_verified, font=('Helvetica', 13, 'bold'), fg='red')
# NameOfUserLabelHospitalpage2.pack();
# NameOfUserLabelHospitalpage2.place(x=490, y=20);


#my_label=Label(root,text='');

#Empty page/table for new order
OrdersTree = ttk.Treeview(ordersFrame,height=20);

# Defining number of columns
OrdersTree["columns"]=("1","2","3","4");

OrdersTree.pack(side=LEFT, padx=100, pady=110)

#Order main page scrollbar-vertical
OrderMainPagescroll = Scrollbar(ordersFrame, orient="vertical", command=OrdersTree.yview);
# OrderMainPagescroll.pack(side=RIGHT,fill=Y);
OrderMainPagescroll.place(x=705, y=110, height=330)

OrdersTree.configure(yscrollcommand = OrderMainPagescroll.set);


#Formate Columns
OrdersTree.column("#0",width=0,minwidth=0);
OrdersTree.column("1");
OrdersTree.column("2",width=120);
OrdersTree.column("3",width=40);
OrdersTree.column("4");


#Define headers/titles in table
OrdersTree.heading("#0", text="Label");
OrdersTree.heading("1", text="Hospital");
OrdersTree.heading("2", text="Injection Date");
OrdersTree.heading("3", text="Doses");
OrdersTree.heading("4", text="Last Updated");

#clear/delete Order main page tree-all records
def clear_tree():
    OrdersTree.delete(*OrdersTree.get_children());


#clear/delete Edit page tree-all records
def ClearEdittree_2():
    #"""Initialization-clear/Delete all the records"""
    for rawselected in EditTree.get_children():
        EditTree.delete(rawselected);

ListofCurrnetHospitalOrderMainPage=[];

global counterOrderId;

#if counterOrderId==0:
counterOrderId=0;

def updateOrdersTreeMainPageOutputOnly():

    clear_tree();
    # # Absorb Orders list data from db
    cursor = db.cursor();

    #Show output to Order main page tree-id,date,sum of doses,and last updated
    #cursor.execute("SELECT hospitalID,Date,SUM(amount) FROM orders GROUP BY Date,hospitalID;");
    cursor.execute("SELECT hospital.Name,orders.Date,SUM(orders.amount),MAX(updated_at) FROM orders INNER JOIN hospital ON hospital.idhospital = orders.hospitalID  GROUP BY orders.Date,orders.hospitalID;");
    SumOFAmount1 = cursor.fetchall();
    print(SumOFAmount1);
    # for x in SumOFAmount1:
    #     print(x);
    #convert list of tuples into list
    # ListofSumOfAmountPerHospital1 = [item1 for t1 in ListofSumOfAmountPerHospital for item1 in t1];

    # #Remove any duplicates from a List:-method 1
    # mylist = list(dict.fromkeys(ListofSumOfAmountPerHospital1))
    # ListOfSumOfAmount = [item1 for t1 in mylist for item1 in t1];

    # remove duplicated from list-method 2
    # FIxedHospitalList = []
    # for i in ListofCurrnetHospitalOrderMainPage:
    #     if i not in FIxedHospitalList:
    #         FIxedHospitalList.append(i)
    #
    # # printing list after removal
    # print ("The list after removing duplicates : " + str(FIxedHospitalList))

#         cursor.execute("INSERT INTO orders (idorders,Date,injection_time,amount,hospitalID,batchID,DecayCorrected) VALUES (%s,%s,%s,%s,%s,%s,%s);",ValuseTuple);
    #output orders main data from DB to the orders tree
    for record in SumOFAmount1:
        OrdersTree.insert(parent='', index='end', text='',values=(record[0],record[1],record[2],record[3]))#record[0]=Idhospital,record[1]=Injection time,record[2]=Amount of doses,record[3]=last updated
        ListofCurrnetHospitalOrderMainPage.append(record[0]);

    #print(ListofCurrnetHospitalOrderMainPage);
    root.wm_state('normal');#Open orders main page
    #OrdersTree.pack();

    db.commit();
    cursor.close();

updateOrdersTreeMainPageOutputOnly();

def SearchOutpout(data):
    """Function for updating List of mian tree order main-output for searching component """

    clear_tree();
    # # Absorb Orders list data from db
    cursor = db.cursor();

    #Show output to Order main page tree-id,date,sum of doses
    #cursor.execute("SELECT hospitalID,Date,SUM(amount) FROM orders GROUP BY Date,hospitalID;");
    cursor.execute(f'SELECT hospital.Name,orders.Date,SUM(orders.amount),MAX(updated_at) FROM orders INNER JOIN hospital ON hospital.idhospital = orders.hospitalID GROUP BY orders.Date,orders.hospitalID;');
    DataAchievedBtSearch = cursor.fetchall();
    for x in DataAchievedBtSearch:
        print(x);
    #convert list of tuples into list
    # y = [item1 for t1 in SumOFAmount1 for item1 in t1];
    # print(y);
    #print(DataAchievedBtSearch);
    countertemp=0;
    for record in DataAchievedBtSearch:
        print(record)
        OrdersTree.insert(parent='', index='end', text='',values=(data[countertemp],record[1],record[2],record[3]));   #record[0]=hospital name,record[1]=Injection time,record[2]=Amount of doses,record[3]=updated_at
        countertemp=countertemp+1;


    root.wm_state('normal');#Open orders main page
    OrdersTree.pack();

    db.commit();
    cursor.close()



#########################Orders main pages buttons###############################

#Create Refresh(DB) button
global reafreahImg;
reafreshIcon = Image.open("./Images/regreshButton.png");
resizedReafreshEditIcon = reafreshIcon.resize((23,23), Image.ANTIALIAS);
reafreahImg = ImageTk.PhotoImage(resizedReafreshEditIcon);
ReafrshButton=Button(ordersFrame, image=reafreahImg, borderwidth=0,command=updateOrdersTreeMainPageOutputOnly)
ReafrshButton.pack();
ReafrshButton.place(x=530, y=63);


#Create Search window
searchEntry = Entry(ordersFrame,font=("Halvetica",12));
searchEntry.insert(0, 'Search Hospital Name');
searchEntry.pack();
searchEntry.place(x=640, y=65);

#Create search icon
searchIcon = Image.open("./Images/SearchButton.png");
resizedSearchedEditIcon = searchIcon.resize((23,23), Image.ANTIALIAS);
SearchImg = ImageTk.PhotoImage(resizedSearchedEditIcon);
SearchLabelicon=Label(image=SearchImg);
SearchLabelicon.pack();
SearchLabelicon.place(x=610, y=135);

def SearchComponent(event):
      """Function for creating search component """

      #get the Orders data from DB list-refresh
      cursor = db.cursor();

    #Show output to Order main page tree-id,date,sum of doses
    #cursor.execute("SELECT hospitalID,Date,SUM(amount) FROM orders GROUP BY Date,hospitalID;");
      cursor.execute("SELECT hospital.Name,orders.Date,SUM(orders.amount),MAX(updated_at) FROM orders INNER JOIN hospital ON hospital.idhospital = orders.hospitalID GROUP BY orders.Date,orders.hospitalID;");
      SumOFAmount1 = cursor.fetchall();
      print(SumOFAmount1);
      ListofCurrnetHospitalOrderMainPage=[];
      for record in SumOFAmount1:
         OrdersTree.insert(parent='', index='end', text='',values=(record[0],record[1],record[2],record[3])) #record[0]=Idhospital,record[1]=Injection time,record[2]=Amount of doses,record[3]=updated_at
         ListofCurrnetHospitalOrderMainPage.append(record[0]);

    #print(ListofCurrnetHospitalOrderMainPage);
      root.wm_state('normal');#Open orders main page
    #OrdersTree.pack();

      db.commit();
      cursor.close();



      typed = searchEntry.get();
      if typed == '':#if nothing typed in the entry
          dataList = ListofCurrnetHospitalOrderMainPage;#return all of curent list(relevant orders)
      else:
          dataList=[];
          for item in ListofCurrnetHospitalOrderMainPage:
              if typed.lower() in item.lower():           #catch capital and lower case
                  dataList.append(item);
                  print("Hospital finded:",dataList);

      SearchOutpout(dataList);#update past list in the new searched list

searchEntry.bind("<KeyRelease>",SearchComponent)#catch any key pressed and released from keyboard- event

#Clear search button on-click in the input widget
# define an event handler
def handleEvent(event):
    searchEntry.delete(0, END);

# bind the click of the Entry to the handler
searchEntry.bind("<ButtonRelease>", handleEvent)


#######Filter by material-drop down##################################################
# MaterialClickDropDownMenu = StringVar();
# MaterialClickDropDownMenu.set("Material"); #default value
#
# Absorb materials list data from db
cursor = db.cursor();
cursor.execute("SELECT idmaterial,materialName FROM material");
Material_in_db = cursor.fetchall();
print(Material_in_db);
AllListOption=("All");
Material_in_db.append(AllListOption);
print(Material_in_db);
def updateOrdersTreeByMaterialFiltering(materialSelData):
    clear_tree();
    # # Absorb Orders list data from db
    cursor = db.cursor();
    if materialSelData=='A':
    #Show output to Order main page tree-id,date,sum of doses
     cursor.execute(f"SELECT hospital.Name,orders.Date,SUM(orders.amount),MAX(updated_at) FROM orders  INNER JOIN hospital ON hospital.idhospital = orders.hospitalID GROUP BY orders.Date,orders.hospitalID ;");
    else:
    #Show output to Order main page tree-id,date,sum of doses filtering by Material ID
     cursor.execute(f"SELECT hospital.Name,orders.Date,SUM(orders.amount),MAX(updated_at) FROM orders  INNER JOIN hospital ON hospital.idhospital = orders.hospitalID WHERE materialID={materialSelData}  GROUP BY orders.Date,orders.hospitalID ;");

    filteringRowsFromDB = cursor.fetchall();
    print(filteringRowsFromDB);
    for record in filteringRowsFromDB:
        OrdersTree.insert(parent='', index='end', text='',values=(record[0],record[1],record[2],record[3]))#record[0]=Idhospital,record[1]=Injection time,record[2]=Amount of doses,record[3]=last updated
        ListofCurrnetHospitalOrderMainPage.append(record[0]);

    #print(ListofCurrnetHospitalOrderMainPage);
    root.wm_state('normal');#Open orders main page
    OrdersTree.pack();

    db.commit();
    cursor.close();


# #function to get the value from the drop down menu
def MaterialsSelectedeFilteringFunc(event):
    #MaterialDropDownLabel=Label(ordersFrame,text=MaterialsDropDownFilteringMainPage.get()).pack()
    MaterialSelected=MaterialsDropDownFilteringMainPage.get();
    Mlist=list(MaterialSelected);
    print(Mlist[0]);
    updateOrdersTreeByMaterialFiltering(Mlist[0]);#Calling to output sql function


MaterialsDropDownFilteringMainPage = ttk.Combobox(ordersFrame,state="readonly",value=Material_in_db,width=9);
MaterialsDropDownFilteringMainPage.current(2);

MaterialsDropDownFilteringMainPage.bind("<<ComboboxSelected>>",MaterialsSelectedeFilteringFunc)
MaterialsDropDownFilteringMainPage.pack();
MaterialsDropDownFilteringMainPage.place(x=370, y=70);

############################################################################

# Absorb hosital list data from db
# cursor = db.cursor();
# cursor.execute("SELECT idhospital,Name FROM hospital");
# hospitals_in_db = cursor.fetchall();
# HospitalListForDeleteOrder = hospitals_in_db;
# print(HospitalListForDeleteOrder);


OrderselectedEvent = tk.StringVar();
def deleteOrderEvent(event):

        """Function for removing order from DB"""
        global DoseNum_Selected,DateSelected,InjectionTimeSelected,IDofHospitalSelected2;
        row = OrdersTree.focus();

        dataofchoosnenRowListEditTree=row;
        #print(dataofchoosnenRowListEditTree);
        DataOfRowSelectedDicEditTree=OrdersTree.item(dataofchoosnenRowListEditTree);
        DataOfRowSelectedList=DataOfRowSelectedDicEditTree['values'];
        print("Record/Order selected to delete: ",DataOfRowSelectedList);
        hospitalSelected=DataOfRowSelectedList[0];
        DateSelected=DataOfRowSelectedList[1];
        InjectionTimeSelected=DataOfRowSelectedList[2];

        #search hospital by name from hospital table db and get the ID as output
        cursor = db.cursor();
        cursor.execute(f'SELECT CAST(idhospital AS SIGNED) FROM hospital WHERE Name="{hospitalSelected}"');
        IDofHospitalSelected1 = cursor.fetchall();
        TempID=[i[0] for i in IDofHospitalSelected1];#find index number in a list of tuple
        IDofHospitalSelected2=int(TempID[0]);
        print(f'{IDofHospitalSelected2} : {hospitalSelected}');

def deleteOrderfunc():
    MsgBox = messagebox.askquestion ('Info message','Are you sure you want to cancel/delete the order?',icon = 'warning')
    if MsgBox == 'yes':
        rawSelectedToDelete=OrdersTree.selection();#selected item:I001,I002,I003....
        RecoredDeletedFlug=1;
        try:
            cursor = db.cursor(buffered=True);
            for rawselected in rawSelectedToDelete:
                UpdateSQlQuery=f"UPDATE  orders SET deleted='{RecoredDeletedFlug}' WHERE  hospitalID= '{IDofHospitalSelected2}' AND Date= '{DateSelected}';";
                DeleteQuery = f"DELETE FROM orders WHERE hospitalID= '{IDofHospitalSelected2}' AND Date= '{DateSelected}';";
                cursor.execute(UpdateSQlQuery);
                cursor.execute(DeleteQuery);
                OrdersTree.delete(rawselected);
                print("DB updated successfully-Record add to deleted column ");
                db.commit();
                cursor.close();
        except mysql.connector.errors.IntegrityError as e:
            messagebox.showinfo("info message","Watch out\nThere is workplan that relate to the order you trying to delete/cancel that still active\n,action cancel!");
            logging.error(traceback.format_exc())
        except Exception as e:
            logging.error(traceback.format_exc())
            print("Error-Order was not updated-please check MySQL");

    else:
        messagebox.showinfo('Return','Return to the Orders screen');



# Remove button (Icon) -Delete Order
global imgDelete;
deleteIcon = Image.open("./‏‏deleteIcon.png")
resizedDeleteIcon = deleteIcon.resize((20,20), Image.ANTIALIAS)
imgDelete = ImageTk.PhotoImage(resizedDeleteIcon)
deleteButton=Button(ordersFrame, image=imgDelete, borderwidth=0,command=deleteOrderfunc)
deleteButton.pack()
deleteButton.place(x=560, y=65)

OrdersTree.bind('<<TreeviewSelect>>', deleteOrderEvent);



###############################################Import File page##################################


def WriteToCsv(result):
    """Function for creating/exporting Excel file"""
    print("try exporting new excel file...");
    headers = ['OrderId', 'Date', 'Injection Time', 'Amount','idhospital','batchID','decayCorrected'];
    with open('orders.csv','a',newline="") as f:
        w = csv.writer(f,dialect='excel');
        messagebox.showinfo("message","Excel file was created");
        # write the headers
        w.writerow(headers);
        for record in result:
            w.writerow(record);


# Absorb Orders table data from db-for excel export
cursor = db.cursor();
cursor.execute("SELECT * FROM orders");
ordersTable_in_db = cursor.fetchall();

#Create Export to Excel buttton
global ExportToCSVImg;
ExportCSVIcon = Image.open("./Images/ExportExcel.png");
resizedExportCSVIcon = ExportCSVIcon.resize((23,23), Image.ANTIALIAS);
ExportToCSVImg = ImageTk.PhotoImage(resizedExportCSVIcon);
ExportToCSVImgicon=Button(ordersFrame, image=ExportToCSVImg, borderwidth=0,command=lambda : WriteToCsv(ordersTable_in_db))
ExportToCSVImgicon.pack();
ExportToCSVImgicon.place(x=585, y=63);

#Create edit icon
# global imgEdit;
# editIcon = Image.open("editIcon.jpg")
# resizedEditIcon = editIcon.resize((20,20), Image.ANTIALIAS)
# imgEdit = ImageTk.PhotoImage(resizedEditIcon)
# editButton=Button(ordersFrame, image=imgEdit, borderwidth=0)
# editButton.pack()
# editButton.place(x=800, y=65)
# edit_button = Button(ordersFrame, text= "Edit")
# edit_button.pack(side= LEFT)
# edit_button.place(x=450, y=50)

#edit field from DB
# query = "UPDATE hospital SET Name = %s ,Fixed_activity_level= %s, Transport_time=%s  WHERE idhospital = %s"
#
# pk = selected_rec[3]
#
# labels = (('Name', ''), ('Fixed activity level', '(mci/h)'),  ('Transport time', '(min)'))
# save_title = "Save Changes"
#
# editHospitalPopup.edit_popup(labels, selected_rec, save_title, query, pk, hospital_tabel)

########################Import File page##################################################################
def importFileFunc():
    AmountListFromDoc=[];
    InjectionTImeListFromdoc=[];
    #ListofVarImportFile=["","","","","",""];
    TempList=["","",""];
    def ImportFilefunction():
        TempNewLISt1=[];
        """This is function for importing Orders files"""
        filename = fd.askopenfilename(
        initialdir="D:\PythonProjects\Cyclotron",
        title="Open a file",
        filetype=(("Word files","*.docx"),("Word files","*.doc"),("Excel files","*.xlsx"),("All Files","*.*"),("PDF files","*.pdf")))
        #print(filename);
        if filename:
            if  "xlsx" in filename :                     #Excel file
                try:
                    filename=r"{}".format(filename)
                    df=pd.read_excel(filename)
                except ValueError:
                    messagebox.showinfo("Error message","File couldn't be open,try again");
                    print("Error");
                except FileNotFoundError:
                    messagebox.showinfo("Error message","File couldn't be open,try again");
                    print("Error");

                #clear_tree();


                #####################################################################
                #Get data from excel
                loc = (str(filename));
                wb = xlrd.open_workbook(loc);
                sheet = wb.sheet_by_index(0);
                sheet.cell_value(0, 0);
                #Get amount data/column from doc/excel-Get amount
                for i in range(1,sheet.nrows):
                    AmountListFromDoc.append(sheet.cell_value(i, 2));
                    print(f"Amount number{i}:{AmountListFromDoc[i-1]}");

                    #Get Injection time data/column from doc/excel-Get Injection time
                for i in range(1,sheet.nrows):
                    InjectionTImeListFromdoc.append(sheet.cell_value(i, 5));
                    print(f'InjectionTIme numner {i}: {InjectionTImeListFromdoc[i-1]}');
                 ########################################################

            #word files

            if "docx" in filename or "doc" in filename:
                #convert word to excel

                if (("doc" in filename) and ("docx" not in filename)):#convert docx to doc
                    doc = aw.Document(filename)
                    filename="NewWordOutput1.docx";
                    doc.save(filename)


                document = Document(filename);
                tables = document.tables;
                df = pd.DataFrame();

                for table in document.tables:
                    for row in table.rows:
                        text = [cell.text for cell in row.cells];
                        df = df.append([text], ignore_index=True);

                #df.columns = ["Column1", "Column2","Column3","Column4","Column5","Column6","Column7","Column8"]
                df.to_excel("D:/PythonProjects/Cyclotron/OrderOutputTest.xlsx");
                #print(df);


                #clear_tree();
                #Get data from excel
                loc = ("D:/PythonProjects/Cyclotron/OrderOutputTest.xlsx");
                wb = xlrd.open_workbook(loc);
                sheet = wb.sheet_by_index(0);
                sheet.cell_value(0, 0);
                #Get amount data/column from doc/excel
                for i in range(1,sheet.nrows):
                    AmountListFromDoc.append(sheet.cell_value(i, 2));
                    print(AmountListFromDoc)

                    #Get Injection time data/column from doc/excel-Get Injection time
                for i in range(1,sheet.nrows):
                    InjectionTImeListFromdoc.append(sheet.cell_value(i, 5));
                    print(f'InjectionTIme numner {i}: {InjectionTImeListFromdoc[i-1]}');
# ##############################################################

        # FileImportedLabel1=Label(ImportFilePage, text="FIle imported successfully",fg="red", font=('Helvetica 12'));
        # FileImportedLabel1.pack();
        # FileImportedLabel1.place(x=450,y=180);

        FileImportedLabel2=Label(ImportFilePage, text=filename, font=('Helvetica 12'),fg="red");
        FileImportedLabel2.pack();
        FileImportedLabel2.place(x=450,y=200);

        root.wm_state('iconic');#minimize orders main page



##########################################
    ImportFilePage = Toplevel(root);
    ImportFilePage.geometry("900x400");
    ImportFilePage.config(bg="#F0F3F4");#color of page-white-gray
    ImportFilePage.title("Import File");
    #NewOrdersecondaryPage = tk.Frame(root);

    ImpoerFilePageLabel=Label(ImportFilePage, text="Import File - Order", font=('Helvetica 17 bold'), fg='#034672');
    ImpoerFilePageLabel.pack();

    HospitalListLabel = Label(ImportFilePage, text="Hospital",bg='white');
    HospitalListLabel.pack();
    HospitalListLabel.place(x=20, y=70);

    #Create hospital drop-down
    # Absorb hosital list data from db
    cursor = db.cursor();
    cursor.execute("SELECT idhospital,Name FROM hospital");
    hospitals_in_db = cursor.fetchall();
    HospitalList2 = hospitals_in_db;


    def HospitalChooseImportFile(HospitalSelectedEvent):
        """Function for create Hospital Drop-Down menu -absorb data from DB"""
        ChoosenHospitalNewOrder=HospitalSelectedImportFile.get();
        print(ChoosenHospitalNewOrder);

        #loop for findinf Id number in the string
        HospitalIDFromChoosenHospital = "";
        for index in ChoosenHospitalNewOrder:
            if index.isdigit():
                HospitalIDFromChoosenHospital = HospitalIDFromChoosenHospital + index;


        TempList[0]=HospitalIDFromChoosenHospital;#HospitalID
        print(TempList[0]);



    HospitalSelectedImportFile = ttk.Combobox(ImportFilePage,state="readonly",value=HospitalList2,width=9);
    HospitalSelectedImportFile.current(0);

    HospitalSelectedImportFile.bind("<<ComboboxSelected>>",HospitalChooseImportFile)
    HospitalSelectedImportFile.pack();
    HospitalSelectedImportFile.place(x=20, y=100);



    # Absorb materials list data from db
    #################Absorb materials list DB#########################
    cursor = db.cursor();
    cursor.execute("SELECT idmaterial,materialName FROM material");
    Material_in_db = cursor.fetchall();
    print(Material_in_db);

    #Create Material Drop-Dowm menu

    #matrial label
    MaterialListLabel = Label(ImportFilePage, text="Material",bg='white');
    MaterialListLabel.pack();
    MaterialListLabel.place(x=20, y=300);

    def MaterialsSelectedeImportFile(MaterialSelectedEvent):
        """Function for create Material Drop-Down menu -absorb data from DB"""
        ChoosenMaterialNewOrder=MaterialsSelectedImportFile.get();
        print(ChoosenMaterialNewOrder);
        print(ChoosenMaterialNewOrder);
        ChoosenMaterial2=ChoosenMaterialNewOrder;
        temp_Var_=list(ChoosenMaterial2);
        TempList[2]=ChoosenMaretrialIDNewOrderManual=int(temp_Var_[0]);
        print(TempList[2]);


    MaterialsSelectedImportFile = ttk.Combobox(ImportFilePage,state="readonly",value=Material_in_db,width=9);
    MaterialsSelectedImportFile.current(0);

    MaterialsSelectedImportFile.bind("<<ComboboxSelected>>",MaterialsSelectedeImportFile)
    MaterialsSelectedImportFile.pack();
    MaterialsSelectedImportFile.place(x=18, y=330);


    def SaveToDB():
        # Function to save order into DB from Import file
        global counterOrderId;
        ######################################################################################################
        def delete_label(label):#Function for Clear label
            label.destroy()


        cursor = db.cursor(buffered=True);
        print("order trying to get in DB-Add pressed");

        try:
           for i in range(1,len(AmountListFromDoc)):
            ValuseTuple=(counterOrderId,i,TempList[1],InjectionTImeListFromdoc[i],AmountListFromDoc[i], TempList[0],TempList[2],0);#BatchId=null
            cursor.execute("INSERT INTO orders (idorders,DoseNumber,Date,injection_time,amount,hospitalID,materialID,DecayCorrected) VALUES (%s,%s,%s,%s,%s,%s,%s,%s);",ValuseTuple);#BatchId=null
            counterOrderId=counterOrderId+1;
           updateOrdersTreeMainPageOutputOnly();##update orders tree main page
        except (mysql.connector.errors.DatabaseError,UnboundLocalError):
            logging.error(traceback.format_exc());
            DateInputCheckMsg=Label(ImportFilePage, text="Please choose date of order",fg="red", font=('Helvetica 12'));
            DateInputCheckMsg.pack();
            DateInputCheckMsg.place(x=20,y=270);
            root.after(5000, delete_label, DateInputCheckMsg); ##Clear label after 5 secondes

            HospitalInputMsg=Label(ImportFilePage, text="Please choose hospital",fg="red", font=('Helvetica 12'));
            HospitalInputMsg.pack();
            HospitalInputMsg.place(x=20,y=130);
            root.after(5000, delete_label, HospitalInputMsg);#Clear label after 5 secondes
        except Exception as e:
            logging.error(traceback.format_exc());
            #messagebox.showerror("Error message","Error !");
            print("Error");

        updateOrdersTreeMainPageOutputOnly();
        ImportFilePage.destroy();##Close import file window
        #Commit changes in DB and close connection
        db.commit()
        cursor.close()


#Create a save button
    # saveFileIcon = Image.open("./Images/saveIcon.png");
    # save_next_Icon = saveFileIcon.resize((100,50), Image.ANTIALIAS);
    # saveImg = ImageTk.PhotoImage(save_next_Icon);
    saveButton=Button(ImportFilePage,text="Save",command=SaveToDB);
    saveButton.pack();
    saveButton.place(x=250, y=320);

    #Create a Cancel button
    # CancelIcon2 = Image.open("./Images/CancelIcon.png");
    # resized_Cancel_Icon2 = CancelIcon2.resize((100,50), Image.ANTIALIAS);
    # CancelImg2 = ImageTk.PhotoImage(resized_Cancel_Icon2);
    CancelButton2=Button(ImportFilePage,text="Cancel",command=lambda: [ImportFilePage.destroy()]);#close window-not working
    CancelButton2.pack();
    CancelButton2.place(x=450, y=320);

    #Create a File button+Label
    FileLabel = Label(ImportFilePage, text="File",bg='white');
    FileLabel.pack();
    FileLabel.place(x=500, y=65);

    FileIcon = Image.open("./Images/FileIcon.png")
    resized_File_Icon = FileIcon.resize((60,60), Image.ANTIALIAS)
    file_Image = ImageTk.PhotoImage(resized_File_Icon)
    FileButton=Button(ImportFilePage, image=file_Image, borderwidth=0,command=ImportFilefunction)
    FileButton.pack()
    FileButton.place(x=500, y=95);

    #Create a Injection Date

    InjectionDateLabel = Label(ImportFilePage, text="Injection Date",bg='white');
    InjectionDateLabel.pack();
    InjectionDateLabel.place(x=20, y=180);
    InjectionDateLabel2 = Label(ImportFilePage, text="Pick a date",fg="gray",font=("Halvetica",10));#fg=color of text
    InjectionDateLabel2.pack();
    InjectionDateLabel2.place(x=20, y=210);

    #Calender
    #add calender icon
    CalendarIcon = Image.open("./Images/CalendarIcon.png");
    resizedCalenderIcon = CalendarIcon.resize((23,23), Image.ANTIALIAS);
    CalenderImg = ImageTk.PhotoImage(resizedCalenderIcon);
    CalenderLabelicon=Label(ImportFilePage,image=CalenderImg);
    CalenderLabelicon.pack();
    CalenderLabelicon.place(x=180, y=238);

    date_today = datetime.now() # today's date

    #Add calender widget/method
    selectedDate=tk.StringVar() # declaring string variable
    def print_sel(e):
        ChoosenDateForImport=cal1.get_date();
        TempList[1]=ChoosenDateForImport;
        print( TempList[1]);

    cal1=DateEntry(ImportFilePage,background='darkblue',foreground='white',mindate=date_today,textvariable=selectedDate);
    cal1.pack(pady = 20);
    cal1.place(x=20, y=240);
    cal1.bind("<<DateEntrySelected>>", print_sel);#catch date event

    #cal1=Calendar(ImportFilePage,font="Arial 14", selectmode='day',year=2022, month=5, disableddaybackground="red", day=22);

    # for i in range(len(page_stuff)):
    #  OrdersTree.insert('', 'end',values=i)
    #OrdersTree.insert(parent='', index=0, text='', values=(TempList[0],TempList[1]));

    # def my_upd(*args): # triggered when value of string varaible changes
    #     l1.config(text=sel.get()); # read and display date
    #
    # l1=tk.Label(ImportFilePage,bg='yellow');  # Label to display date
    # l1.pack();
    #
    # sel.trace('w',my_upd) ;# on change of string variable
    # print(sel);

    #print(TempList);
    ImportFilePage.mainloop();



# iid=0
# for hospital in hospitals_in_db:
#     #print(hospital)
#     hospitals_list.insert(parent='', index='end', iid=iid, text='',
#                           values=(hospital[1], hospital[2], hospital[3]))
#     iid +=1
#
# hospitals_list.pack()


###################Buttons for edit,delete,import file and etc.###################################
#Create a button in the main Window to open the popup
# edit_button = Button(hospitalFrame, text= "Edit", command= open_popup_hospital)
# edit_button.pack(side= LEFT)
# edit_button.place(x=450, y=50)
# edit_button.pack(side=LEFT, padx=PlaceLable_X+100, pady=PlaceLable_Y+50)

#clear/delete Edit page tree-all records
# def disabledTree():
#     #"""Initialization-clear/Delete all the records"""
#     for child in NewOrderTree_P2.winfo_children():
#         child.configure(state='disable')

############################new order page###################################################

def PopUpForNewOrder():

    root = tk.Tk()
    root.title("New Order")
    root.geometry("700x600");
    root.configure(bg="#F0F3F4")

    #NewOrderMainPage =Toplevel(root);
    NewOrderMainPage =tk.Frame(root);
    NewOrderMainPage.config(bg="#F0F3F4");#color of page-white-gray


    #Create secnd pop-up window for order page
    NewOrdersecondaryPage = tk.Frame(root);
    NewOrdersecondaryPage.config(bg="#F0F3F4");#color of page-white-gray


########################page number 1,New order page#########################################################

##

    NeworderTitleLabel=Label(NewOrderMainPage, text="New Order ",bg="#F0F3F4", font=('Helvetica 17 bold'), fg='#034672');
    NeworderTitleLabel.pack();
    NeworderTitleLabel.place(x=270,y=25);
    # labels
    #Create hospital Drop-down menu

    HospitalListLabel = Label(NewOrderMainPage, text="Hospital",bg='white');
    HospitalListLabel.pack();
    HospitalListLabel.place(x=20, y=70);

    # Absorb hosital list data from db
    cursor = db.cursor();
    cursor.execute("SELECT idhospital,Name FROM hospital");
    hospitalsListForNewOrderManual = cursor.fetchall();

    def HospitalChooseNewOrder(HospitalSelectedEvent):
        """Function for create Hospital Drop-Down menu -absorb data from DB"""
        global hospitalId;
        ChoosenHospitalNewOrder=(HospitalSelectedNewOrder.get());

        #loop for find hospital name in the string
        HospitalNameFromChoosenHospital = "";
        for index in ChoosenHospitalNewOrder:
            if index.isdigit():
             pass;
            else:
             HospitalNameFromChoosenHospital = HospitalNameFromChoosenHospital + index;


        ChoosenhospitalNameFromDropDown=HospitalNameFromChoosenHospital;

        print(ChoosenhospitalNameFromDropDown);
        #Print to the screen the hospital selected-print to page number 2
        HospitalLabel2=Label(NewOrdersecondaryPage, text= ChoosenhospitalNameFromDropDown, bg="white", font=('Helvetica 14'));
        HospitalLabel2.pack();
        HospitalLabel2.place(x=20,y=80);

        #loop for findinf Id number in the string
        HospitalIDFromChoosenHospital = "";
        for index in ChoosenHospitalNewOrder:
            if index.isdigit():
                HospitalIDFromChoosenHospital = HospitalIDFromChoosenHospital + index;


        hospitalId=int(HospitalIDFromChoosenHospital);


    HospitalSelectedNewOrder = ttk.Combobox(NewOrderMainPage,state="readonly",value=hospitalsListForNewOrderManual,width=15);
    HospitalSelectedNewOrder.current(0);

    HospitalSelectedNewOrder.bind("<<ComboboxSelected>>",HospitalChooseNewOrder);
    HospitalSelectedNewOrder.pack();
    HospitalSelectedNewOrder.place(x=20, y=100);


###########################################################################

    # Absorb materials list data from db
    cursor = db.cursor();
    cursor.execute("SELECT idmaterial,materialName FROM material");
    Material_in_db1 = cursor.fetchall();
    print(Material_in_db1);

    #Create Material Drop-Dowm menu

    #matrial label
    MaterialListLabel1 = Label(NewOrderMainPage, text="Material",bg='white');
    MaterialListLabel1.pack();
    MaterialListLabel1.place(x=400, y=310);


    def MaterialsSelectedeNewOrder(MaterialSelectedEvent):
        global ChoosenMaretrialIDNewOrderManual;
        """Function for create Material Drop-Down menu -absorb data from DB"""
        ChoosenMaterialNewOrder=MaterialsSelectedNeworder.get();
        print(ChoosenMaterialNewOrder);
        temp_Var_=list(ChoosenMaterialNewOrder);
        ListofVal[7]=ChoosenMaretrialIDNewOrderManual=int(temp_Var_[0]);
        print(ListofVal[7]);


    MaterialsSelectedNeworder = ttk.Combobox(NewOrderMainPage,state="readonly",value=Material_in_db1,width=9);
    MaterialsSelectedNeworder.current(0);

    MaterialsSelectedNeworder.bind("<<ComboboxSelected>>",MaterialsSelectedeNewOrder);

    MaterialsSelectedNeworder.pack();
    MaterialsSelectedNeworder.place(x=400, y=330);
    global OrderID,idCounter,amount,ChoosenMaretrialIDNewOrderManual;
    OrderID=0;
    # declaring string variable for storing time interval
    TimeIntervals=tk.StringVar(NewOrderMainPage);
    # declaring string variable for storing time
    HoursVar = StringVar(NewOrderMainPage);
    MinutesVar = StringVar(NewOrderMainPage);
    amountVar=tk.StringVar(NewOrderMainPage);

    # declaring string variable for storing amount
    ListofVal=["","","","","","","",""];

    ListofTimeIntervals=[];


    def submitToNextPage():
        global amount;
        global hospitalId;
        global OrderID;
        #disabledTree();
        #Initialization-clear/Delete all the records
        for rawselected in NewOrderTree_P2.get_children():
            NewOrderTree_P2.delete(rawselected);

        def destroy_widget(widget):
            widget.destroy()


        #Get event values-events
        Time_Intervals=TimerangeLabelEntry.get();
        IntAmount=AmountOfDosesLabelEntry.get();
        Minutes_Var=MinutesCLockSelected.get();
        Hours_Var=HoursClockedSelected.get();

        #message box if not try to click next if inputs are empty-Input check
        #Check input material
        # try:
        #     ChoosenMaretrialIDNewOrderManual;
        # except NameError:
        #     MatrerialInputCheckMsg=Label(NewOrderMainPage, text="Please choose Material ",fg="red", font=('Helvetica 12'));
        #     MatrerialInputCheckMsg.pack();
        #     MatrerialInputCheckMsg.place(x=400,y=360);
        #     root.after(5000, destroy_widget, MatrerialInputCheckMsg) ##Clear label after 5 secondes
        # else:
        #     print(ChoosenMaretrialIDNewOrderManual);
        #try:
        cursor = db.cursor();
        SearchSpecOrderQueryByDoubleclick=f'SELECT idorders,CONCAT(hospitalID,Date) AS AdressPlusHospital FROM orders where hospitalID="{hospitalId}" AND Date="{ChoosenDateForManaulOrder}"';
        cursor.execute(SearchSpecOrderQueryByDoubleclick);
        OrderDatatoSpecificOrder = cursor.fetchall();

        db.commit();
        cursor.close();

        for x in OrderDatatoSpecificOrder:
            print(f"order selected:{x}");

        if OrderDatatoSpecificOrder:
            messagebox.showerror("Error message","There is allready order with that date and hospital,cant continue!");
            raise Exception("There is allready order with that date and hospital,cant continue");


        try:
            if ChoosenMaretrialIDNewOrderManual in globals():
                MaterialInputCheckMsg=Label(NewOrderMainPage, text="Please  choose Material ",fg="red", font=('Helvetica 12'));
                MaterialInputCheckMsg.pack();
                MaterialInputCheckMsg.place(x=400,y=360);
                print("bla bla")
                root.after(5000, destroy_widget, MaterialInputCheckMsg);#Clear label after 5 secondes
            else:
                print(ChoosenMaretrialIDNewOrderManual);

        except NameError:
            MatrerialInputCheckMsg=Label(NewOrderMainPage, text="Please choose Material ",fg="red", font=('Helvetica 12'));
            MatrerialInputCheckMsg.pack();
            MatrerialInputCheckMsg.place(x=400,y=360);
            root.after(5000, destroy_widget, MatrerialInputCheckMsg) ##Clear label after 5 secondes
        else:
            print(ChoosenMaretrialIDNewOrderManual);

        try:
            ChoosenDateForManaulOrder;
        except NameError:
            DateInputCheckMsg=Label(NewOrderMainPage, text="Please choose date of order",fg="red", font=('Helvetica 12'));
            DateInputCheckMsg.pack();
            DateInputCheckMsg.place(x=20,y=260);
            root.after(5000, destroy_widget, DateInputCheckMsg) ##Clear label after 5 secondes
        else:
            print(ChoosenDateForManaulOrder);

        try:
         hospitalId
        except NameError:
         HospitalInputMsg=Label(NewOrderMainPage, text="Please choose hospital",fg="red", font=('Helvetica 12'));
         HospitalInputMsg.pack();
         HospitalInputMsg.place(x=20,y=140);
         root.after(5000, destroy_widget, HospitalInputMsg);#Clear label after 5 secondes
        else:
         print(hospitalId);

         #loop for findinf if number in the string-time intervals
        for index in Time_Intervals:
            if index.isdigit()==False:
               NaNFlag=True;
            else:
                NaNFlag=False;

        if (Time_Intervals=='') or (NaNFlag==True):
            TImeIntervalMsg=Label(NewOrderMainPage, text="Please fill in Time Intervals",fg="red", font=('Helvetica 12'));
            TImeIntervalMsg.pack();
            TImeIntervalMsg.place(x=400,y=270);
            root.after(5000, destroy_widget, TImeIntervalMsg);#Clear label after 5 secondes
        else:
            print(Time_Intervals);

        if (Minutes_Var=='0') or (Hours_Var=='0'):
            BegTImeMsg=Label(NewOrderMainPage, text="Please fill in Beginning Time ",fg="red", font=('Helvetica 12'));
            BegTImeMsg.pack();
            BegTImeMsg.place(x=20,y=370);
            root.after(5000, destroy_widget, BegTImeMsg);#Clear label after 5 secondes
        else:
            print(Minutes_Var,Hours_Var);


        for index in IntAmount:
         if index.isdigit()==False:
             AmountOfrowsMsg1=Label(NewOrderMainPage, text="Amount of rows can accept only integer numbers ",fg="red", font=('Helvetica 12'));
             AmountOfrowsMsg1.pack();
             AmountOfrowsMsg1.place(x=400,y=165);
             root.after(5000, destroy_widget, AmountOfrowsMsg1);#Clear label after 5 secondes
        # else:
         #    NaNFlag=False;

        try:
            IntAmount=int(IntAmount);
        except ValueError:
            AmountOfrowsMsg2=Label(NewOrderMainPage, text="Please fill in Amount-Of-Rows for order ",fg="red", font=('Helvetica 12'));
            AmountOfrowsMsg2.pack();
            AmountOfrowsMsg2.place(x=400,y=140);
            root.after(5000, destroy_widget, AmountOfrowsMsg2);#Clear label after 5 secondes
        else:
            print(IntAmount);


        #print("Error")
        ListofVal[0]=idCounter=1;
        ListofVal[1]=(IntAmount/IntAmount);
        ListofVal[2]=int(Hours_Var);
        ListofVal[3]=int(Minutes_Var);
        TimeInjectionVar=f'{ListofVal[2]}:{ListofVal[3]}';
        print(TimeInjectionVar);
        ListofVal[4]=IntAmount;
        ListofVal[5]=hospitalId;
        ListofVal[6]=int(Time_Intervals);



        for record in range(int(IntAmount)):
            NewOrderTree_P2.insert("", "end",values=( ListofVal[0],ListofVal[1],f'{ListofVal[2]}:{ListofVal[3]}'));#ListofVal[0]=id,ListofVal[1]=amount/amount,ListofVal[2]=hours:ListofVal[3]=minutes
            #ListofVal[2]=ListofVal[2]+1;       #Hours jumps/intervals
            ListofVal[3]=ListofVal[6];         #Add minutes intervals
            ListofVal[0]= ListofVal[0]+1;
            TimeInjectionVar=f'{ListofVal[2]}:{ListofVal[3]}';
            print("Time per Dose:",TimeInjectionVar);
            ListofTimeIntervals.append(TimeInjectionVar);
            ListofVal[2]=ListofVal[2]+1;       #Hours jumps/intervals

        OrderID+=1;#counterID=counterID+1
        #print(OrderID);
        amountVar.set("");
        MinutesVar.set("");
        HoursVar.set("");
        print("List of intervals:",ListofTimeIntervals);

        deleteButton2['state'] = DISABLED;#Disable remove button
        AddRowButton['state'] = DISABLED;#Disable add dose/record button
        NewOrderTree_P2.state(('disabled',));#Disabled/gray-out new order tree
        dropDownInjectionT_M.configure(state="disabled");#Disabled/gray-out time-injection drop-down menu
        dropDownAmountM.configure(state="disabled");#Disabled/gray-out change amount drop-down menu
        #swap function for viewing New order page,frame 2,after pressing "next" """
        NewOrderMainPage.forget();
        NewOrdersecondaryPage.pack(fill='both',expand=1);

    # def clearwidgets():
    #     AmountOfrowsMsg1.destroy(),AmountOfrowsMsg2.destroy(),BegTImeMsg.destroy();
    #     TImeIntervalMsg.destroy(),HospitalInputMsg.destroy(),DateInputCheckMsg.destroy();
    #
    #
    #
    # searchEntry.bind("<Enter>",SearchComponent)#catch any key pressed and released from keyboard- event

    #Create Amount of Doses input
    AmountOfDosesLabel = Label(NewOrderMainPage, text="Amount of rows",bg="white");
    AmountOfDosesLabel.place(x=400, y=80);

    AmountOfDosesLabelEntry = Entry(NewOrderMainPage,textvariable=amountVar,font=("Halvetica",12));
    AmountOfDosesLabelEntry.config(width=7);#width of window

   # amount=AmountOfDosesLabelEntry.get();
    # print(amount);
    AmountOfDosesLabelEntry.insert(0, '');
    AmountOfDosesLabelEntry.pack();
    AmountOfDosesLabelEntry.place(x=400, y=120);


    #Create Injection time input/entry
    InjectionTimeLabel = Label(NewOrderMainPage, text="Injection Date",bg="white");
    InjectionTimeLabel.place(x=20, y=200);

    #Calender
    #add calender icon
    global  CalenderImg1
    CalendarIcon1 = Image.open("./Images/CalendarIcon.png");
    resizedCalenderIcon1 = CalendarIcon1.resize((23,23), Image.ANTIALIAS);
    CalenderImg1 = ImageTk.PhotoImage(resizedCalenderIcon1,master=NewOrderMainPage);
    CalenderLabelicon1=Label(NewOrderMainPage,image=CalenderImg1);
    CalenderLabelicon1.pack();
    CalenderLabelicon1.place(x=172, y=237);

    date_today = datetime.now() # today's date

    #Add calender widget/method
    selectDateEventManaulOrder=tk.StringVar(NewOrdersecondaryPage) # declaring string variable
    selectDateEventManaulOrder.set(' ');
    def print_sel(e):
        global ChoosenDateForManaulOrder;
        """ This function print to the tree/table """
        ChoosenDateForManaulOrder=cal.get_date();
        print(ChoosenDateForManaulOrder);
        #copy and past date event to page number 2
        dateLabel2=Label(NewOrdersecondaryPage, text= ChoosenDateForManaulOrder, bg="white", font=('Helvetica 12'));
        dateLabel2.pack();
        dateLabel2.place(x=20,y=110);

    def destroyNewOrderFunc():
        """Function for cancel button"""
        root.destroy();
        updateOrdersTreeMainPageOutputOnly();#Refresh/Update Main page
        OrdersTree.pack();         #open order main page immedaitly

    cal=DateEntry(NewOrderMainPage,background='darkblue',mindate=date_today,foreground='white',textvariable=selectDateEventManaulOrder);
    cal.pack(pady = 20);
    cal.config(width=20);#width of window
    cal.place(x=20, y=240);
    cal.bind("<<DateEntrySelected>>", print_sel);#catch date event


    #
    #Create Time range input/Entry
    TimerangeLabel = Label(NewOrderMainPage, text="Time Range/Intervals",bg="white");
    TimerangeLabel.place(x=400, y=200);

    TimerangeLabelEntry = Entry(NewOrderMainPage,textvariable=TimeIntervals,font=("Halvetica",12));
    TimerangeLabelEntry.config(width=7);#width of window
    #TimerangeLabelEntry.insert(0, '');
    TimerangeLabelEntry.pack();
    TimerangeLabelEntry.place(x=400, y=240);

    #Create Beginng time input
    AmountOfDosesLabel = Label(NewOrderMainPage, text="Beginning time",bg="white");
    AmountOfDosesLabel.place(x=20, y=300);
    # AmountOfDosesLabelEntry = Entry(NewOrderMainPage,font=("Halvetica",12));
    # AmountOfDosesLabelEntry.insert(0, '');
    # AmountOfDosesLabelEntry.pack();
    # AmountOfDosesLabelEntry.place(x=20, y=270);
    InjectionDateLabel2 = Label(NewOrderMainPage, text="Pick a time",fg="gray",font=("Halvetica",10));#fg=color of text
    InjectionDateLabel2.pack();
    InjectionDateLabel2.place(x=20, y=330);

    ################FIxed hour,minutes,and secondes########################
    #Hours
    HoursClockedSelected = Spinbox(NewOrderMainPage,textvariable=HoursVar,from_= 0, to = 24,wrap=True,width=2);
    HoursClockedSelected.pack();
    HoursClockedSelected.place(x=20, y=350);
    # Minutes
    MinutesCLockSelected = Spinbox(NewOrderMainPage,textvariable=MinutesVar ,from_= 0, to = 59,wrap=True,width=2);
    MinutesCLockSelected.pack();
    MinutesCLockSelected.place(x=65, y=350);

    # #Create a Cancel button
    # global CancelImg;
    # CancelIcon = Image.open("./Images/CancelButton.png");
    # resized_Cancel_Icon = CancelIcon.resize((100,50), Image.ANTIALIAS);
    # CancelImg = ImageTk.PhotoImage(resized_Cancel_Icon);
    CancelButtonNewOrderPage1=Button(NewOrderMainPage,text="Cancel",command=destroyNewOrderFunc);#close window-not working
    CancelButtonNewOrderPage1.pack();
    CancelButtonNewOrderPage1.place(x=400, y=530);


    #Submit/Next button
    #global NextButton;
    # nextIcon = Image.open("./Images/nextButton.png");
    # resized_next_Icon = nextIcon.resize((100,50), Image.ANTIALIAS);
    # NextButton = ImageTk.PhotoImage(resized_next_Icon);
    sub_btn=tk.Button(NewOrderMainPage ,text="Next", command = submitToNextPage)
    sub_btn.pack();
    sub_btn.place(x=200, y=530)

    def submitToNextPage1(e):
        """Enter key pressed replace next button"""
        submitToNextPage();
    root.bind('<Return>', submitToNextPage1)


    NewOrderMainPage.pack(fill='both',expand=1);


    #####################################New order 2, second page numer 2###################################################

    NewOrdersecondaryLabel=Label(NewOrdersecondaryPage, text="New Order",bg="#F0F3F4", font=('Helvetica 17 bold'), fg='#034672');
    NewOrdersecondaryLabel.pack();
    NewOrdersecondaryLabel.place(x=270,y=27);


    def enterToDB():#Function to insert data into My-SQL Db
        # ordersFrame.forget();
        # toolbar.forget();
        #ordersFrame.wm_state('iconic');#minimize orders main page
        global counterOrderId;
        MsgBox = messagebox.askquestion ('Info message','Do you wish to proceed? every changed will be saved in the DB',icon = 'warning')
        if MsgBox == 'yes':
            # ValuseDic = {
            #        'idorders': 4,
            #        'Date': '2002-03-92',
            #        'injection_time': '11:20:11',
            #        'amount': 7,
            #        'hospitalID': 7,
            #        'batchID': 7,
            #        'DecayCorrected': 7 }  ;

            #Check if all input was enterd
            if (ListofVal[7]!='') and (ListofVal[6]!='') and (ListofVal[5]!='') and (ListofVal[4]!='') and (ListofVal[3]!='') and (ListofVal[2]!='') and (ListofVal[1]!='') and (ListofVal[0]!=''):
                NewOrderTree_P2.state(('!disabled',));#Enable tree items
                deleteButton2['state'] = NORMAL;#Enable remove button
                AddRowButton['state'] = NORMAL;#Enable add dose/record button
                dropDownInjectionT_M.configure(state="normal");#Disabled/gray-out time-injection drop-down menu
                dropDownAmountM.configure(state="normal");#Disabled/gray-out change amount drop-down menu
            else:
                NewOrderTree_P2.state(('disabled',));#Disable tree items
                messagebox.showerror("Error message","Error ! one of the records not filled,please fill order again");
                root.destroy();#close New order page 2
                PopUpForNewOrder();#return to New order page


            if (Stage1forNewOrderBut['state'] == NORMAL):
                SaveOrderbutton['state'] = NORMAL;
                CancelButtonPagenumber2NewORder['state']=DISABLED;
                Stage1forNewOrderBut['state']=DISABLED;
            else:
                SaveOrderbutton['state'] = DISABLED;

            cursor = db.cursor(buffered=True);
            for i in range(1,ListofVal[4]+1):
                ValuseTuple=(counterOrderId,i, ChoosenDateForManaulOrder, ListofTimeIntervals[i-1], ListofVal[1], ListofVal[5],ListofVal[7],0);#BatchId=null
                print("order trying to get in DB-Add pressed");
                try:
                    UpdateSQlQuery="INSERT INTO orders (idorders,DoseNumber,Date,injection_time,amount,hospitalID,materialID,DecayCorrected) VALUES (%s,%s,%s,%s,%s,%s,%s,%s);";#BatchId=null
                    cursor.execute(UpdateSQlQuery,ValuseTuple);
                    counterOrderId=counterOrderId+1;
                    print("DB updated successfully ");
                except Exception as e:
                    logging.error(traceback.format_exc());
                    messagebox.showerror("Error message","Error !");
                    print("Error-Order was not updated-please check MySQL")

            #destroyNewOrderFunc();
            #Commit changes in DB
            db.commit()
            cursor.close()
            #Close connection to DB
            #db.close()
        else:
            messagebox.showinfo('Return','You will now return to the application screen');
            NewOrdersecondaryPage.destroy();#close New order page 2
            PopUpForNewOrder();#return to New order page




    #Create Save order/ADD button
    # global addImg;
    # AddFileIcon = Image.open("./Images/AddButton.png");
    # resized_add_Icon = AddFileIcon.resize((100,50), Image.ANTIALIAS);
    # addImg = ImageTk.PhotoImage(resized_add_Icon);
    Stage1forNewOrderBut=Button(NewOrdersecondaryPage,text="Create Order & Start edit",command=enterToDB);
    Stage1forNewOrderBut.pack();
    Stage1forNewOrderBut.place(x=170, y=520);



    # #Create a Cancel button
    # global CancelImg;
    # CancelIcon = Image.open("./Images/CancelButton.png");
    # resized_Cancel_Icon = CancelIcon.resize((100,50), Image.ANTIALIAS);
    # CancelImg = ImageTk.PhotoImage(resized_Cancel_Icon);
    CancelButtonPagenumber2NewORder=Button(NewOrdersecondaryPage,text="Cancel",command=destroyNewOrderFunc);#close window-not working
    CancelButtonPagenumber2NewORder.pack();
    CancelButtonPagenumber2NewORder.place(x=450, y=520);


    #Empty page/table for new order,create New tree for page 2
    NewOrderTree_P2 = ttk.Treeview(NewOrdersecondaryPage,height=15);
    NewOrderTree_P2.pack();
    NewOrderTree_P2.place(x=170,y=130);

    #New order page 2 page/tree scrollbar -vertical
    NewOrderTree_P2_Scroollbar = Scrollbar(NewOrdersecondaryPage, orient="vertical", command=NewOrderTree_P2.yview);
    NewOrderTree_P2_Scroollbar.place(x=495, y=130, height=330)
    NewOrderTree_P2.configure(yscrollcommand = NewOrderTree_P2_Scroollbar.set);

    # Defining number of columns
    NewOrderTree_P2['columns']= ("ID","Amount","Injection time");
    #Foramte Columns
    NewOrderTree_P2.column("#0",width=0,minwidth=0);
    NewOrderTree_P2.column("ID",anchor=W,width=80,minwidth=25);
    NewOrderTree_P2.column("Amount",anchor=CENTER,width=120,minwidth=25);
    NewOrderTree_P2.column("Injection time",anchor=W,width=120,minwidth=25);

    #Define headers/titles in table
    NewOrderTree_P2.heading("#0", text="Label",anchor=W);
    NewOrderTree_P2.heading("ID", text="Dose Number",anchor=W);
    NewOrderTree_P2.heading("Amount", text="Amount",anchor=CENTER);
    NewOrderTree_P2.heading("Injection time", text="Injection time",anchor=W);

    #############################Amount quantity Event ######################

    AmountList = [
        "0","1","2","3","4","5"
    ]

    status = tk.StringVar()
    status.set("0")

    #Catch event
    def AmountSelectedForNewOrder(event):
        row = NewOrderTree_P2.focus();
        dataofchoosnenRowListEditTree=row;
        print(dataofchoosnenRowListEditTree);
        DataOfRowSelectedDicEditTree=NewOrderTree_P2.item(dataofchoosnenRowListEditTree);
        DataOfRowSelectedList=DataOfRowSelectedDicEditTree['values'];
        print("Raw selectes to edit:",DataOfRowSelectedList);
        IidSelected=DataOfRowSelectedList[0];
        AmountSelected=DataOfRowSelectedList[1];
        InjectionTimeSelected=DataOfRowSelectedList[2];

        if row:
            status.set(NewOrderTree_P2.set(row, 'two'))

    NewOrderTree_P2.bind('<<TreeviewSelect>>', AmountSelectedForNewOrder)

    def set_amountValueEventNewOrderP(CurAmountvalue):
        row = NewOrderTree_P2.focus();
        print("Amount selected",CurAmountvalue);
        if row:
            try:
                cursor = db.cursor(buffered=True);
                UpdateSQlQuery=f"UPDATE  orders SET amount='{CurAmountvalue}'  WHERE DoseNumber = '{DoseNum_Selected}' AND hospitalID= '{ListofVal[5]}' AND Date= '{ChoosenDateForManaulOrder}';";
                cursor.execute(UpdateSQlQuery);
                print("DB updated(Amount) successfully ");
                db.commit();
                cursor.close();
            except Exception as e:
                logging.error(traceback.format_exc())
                print("Error-Order was not updated-please check MySQL")
            NewOrderTree_P2.set(row, '1', CurAmountvalue)


    dropDownAmountM = ttk.OptionMenu(NewOrdersecondaryPage, status, "0", *AmountList, command=set_amountValueEventNewOrderP);
    dropDownAmountM.pack();
    dropDownAmountM.place(x=250, y=490);
    #Change Amount time manual label
    ChangeAmountLabel=Label(NewOrdersecondaryPage, text="Change Amount : ", font=('Helvetica 12'));
    ChangeAmountLabel.pack();
    ChangeAmountLabel.place(x=100,y=490);

    #############################Amount Event over######################

    # ############################Injection Time event#########################
    TimeList = [
        "06:00","06:30","07:00","07:30","08:00","08:30","09:00","09:30","10:00","10:30","11:00",
        "11:30","12:00","12:30","13:00","13:30","14:00","14:30","15:00","15:30","16:00","16:30",
        "17:00","17:30","18:00","18:30","19:00","19:30","20:00"
    ]

    status = tk.StringVar()
    status.set("00:00")

    #Catch Injection  event
    def InjectionTimeselect(event):
        row = NewOrderTree_P2.focus();
        global DoseNum_Selected,AmountSelected,InjectionTimeSelected,DosesSelectedEvent;
        row = NewOrderTree_P2.focus();

        dataofchoosnenRowListEditTree=row;
        print(dataofchoosnenRowListEditTree);
        DataOfRowSelectedDicEditTree=NewOrderTree_P2.item(dataofchoosnenRowListEditTree);
        DataOfRowSelectedList=DataOfRowSelectedDicEditTree['values'];
        print("Raw selectes to edit 2:",DataOfRowSelectedList);
        DoseNum_Selected=DataOfRowSelectedList[0];
        AmountSelected=DataOfRowSelectedList[1];
        InjectionTimeSelected=DataOfRowSelectedList[2];
        if row:
            status.set(NewOrderTree_P2.set(row, 'Injection time'))

    NewOrderTree_P2.bind('<<TreeviewSelect>>', InjectionTimeselect)

    def setInjectionTime(CurTimevalueEditNewOrderP):
        row = NewOrderTree_P2.focus()
        if row:
            try:
                cursor = db.cursor(buffered=True);
                UpdateSQlQuery=f"UPDATE  orders SET injection_time='{CurTimevalueEditNewOrderP}'  WHERE DoseNumber = '{DoseNum_Selected}' AND hospitalID= '{ListofVal[5]}' AND Date= '{ChoosenDateForManaulOrder}';";
                cursor.execute(UpdateSQlQuery);
                print("DB updated(Injection_Time) successfully ");
                db.commit();
                cursor.close();
            except Exception as e:
                logging.error(traceback.format_exc())
                print("Error-Order was not updated-please check MySQL")
            NewOrderTree_P2.set(row, '2', CurTimevalueEditNewOrderP)


    dropDownInjectionT_M = ttk.OptionMenu(NewOrdersecondaryPage, status, "00:00", *TimeList, command=setInjectionTime);
    dropDownInjectionT_M.pack();
    dropDownInjectionT_M.place(x=300, y=460);
    #Change injecion time manual label
    ChangeTImeIjectionLabel=Label(NewOrdersecondaryPage, text="Change Time-Injection : ", font=('Helvetica 12'));
    ChangeTImeIjectionLabel.pack();
    ChangeTImeIjectionLabel.place(x=100,y=460);


    def enterToDBAfterEdit():
        #Function to insert data into My-SQL Db
       destroyNewOrderFunc();
        # root.destroy();#Close import file-manual window
       updateOrdersTreeMainPageOutputOnly();#Refresh/Update Main page
        # OrdersTree.pack();         #open order main page immedaitly




    # #Create Save order/ADD button
    # global addImg;
    # AddFileIcon = Image.open("./Images/AddButton.png");
    # resized_add_Icon = AddFileIcon.resize((100,50), Image.ANTIALIAS);
    # addImg = ImageTk.PhotoImage(resized_add_Icon);
    SaveOrderbutton=Button(NewOrdersecondaryPage,text="Save & Finish Editing",command=enterToDBAfterEdit,state = DISABLED);
    SaveOrderbutton.pack();
    SaveOrderbutton.place(x=320, y=520);


    #Create ADD row button+icon
    # defining a function that will
    # print them on the screen
    #rowTree = StringVar();
    def addRowFunc():
        #global idCounter;
        rowTreetoAdd=(ListofVal[0],ListofVal[1],ListofVal[2]);
        NewOrderTree_P2.insert("", "end", values=rowTreetoAdd);
        ListofVal[0]=ListofVal[0]+1;
        ListofVal[4]+=1;#current amount= courrent amount+1

    def removeRawFunc():
        #rowTree=rowTree.get();
        #for i,j in zip(range(IntAmount),range(BeginigHour,IntAmount)):
        rawSelectedToDelete=NewOrderTree_P2.selection();
        for rawselected in rawSelectedToDelete:
            NewOrderTree_P2.delete(rawselected);
        ListofVal[4]=ListofVal[4]-1;#current amount= current amount-1
    #amountVar.set("");

    #Remove button (Icon) - List-ORders page number 2
    global imgDelete2;
    deleteIcon2 = Image.open("./‏‏deleteIcon.png");
    resizedDeleteIcon2 = deleteIcon2.resize((25,25), Image.ANTIALIAS);
    imgDelete2 = ImageTk.PhotoImage(resizedDeleteIcon2,master=NewOrdersecondaryPage);
    deleteButton2=Button(NewOrdersecondaryPage,image=imgDelete2,bg="white",font=('Helvetica 14'), borderwidth=0,command=removeRawFunc);
    deleteButton2.pack();
    deleteButton2.place(x=465, y=95);




    ####################Buttons for new order-manual page##########################


    global addROWImg;
    AddrowLabel=Label(NewOrdersecondaryPage, text="Add row",bg="white", font=('Helvetica 14'));
    AddrowLabel.pack();
    AddrowLabel.place(x=270,y=98);
    #Add row image+button
    AddrowIcon = Image.open("./addIcon.png");
    resized_add_Row = AddrowIcon.resize((25,25), Image.ANTIALIAS);
    addROWImg = ImageTk.PhotoImage(resized_add_Row,master=NewOrdersecondaryPage);
    AddRowButton=Button(NewOrdersecondaryPage,image=addROWImg, borderwidth=0,command=addRowFunc);
    AddRowButton.pack();
    AddRowButton.place(x=240, y=100);

# ####################end of page number 2 -New order #######################################################################

################################Edit/Update Order main page###################################################################
def UpdateOrder(event):
    global hospitalId;
    curItem = OrdersTree.focus();
    DataOfRowSelectedDic=OrdersTree.item(curItem);
    DataOfRowSelectedList=DataOfRowSelectedDic['values'];
    #print(DataOfRowSelectedList);
    HospitalSelectedName=DataOfRowSelectedList[0];
    DateSelected=DataOfRowSelectedList[1];
    AmountOfDosesSelected=DataOfRowSelectedList[2];



    #search hospital by name from hospital table db and get the ID and the name as output
    cursor = db.cursor();
    cursor.execute(f'SELECT idhospital,Name FROM hospital where Name="{HospitalSelectedName}"');
    hospitalsListForNewOrderManual = cursor.fetchall();

    HospitalListNewOrderPage = hospitalsListForNewOrderManual;
    print(HospitalListNewOrderPage);


    EditPage =Toplevel(root);
    EditPage.title("Edit Order");
    EditPage.geometry("700x600");
    EditPage.config(bg="#F0F3F4");#Color of page(White-Gray)

    #NewOrderMainPage.place(x=450,y=70);

    #NewOrdersecondaryPage = tk.Frame(root);


    HospitalLabelForEditPage=Label(EditPage, text=HospitalSelectedName,bg="white", font=('Helvetica 14'));
    x=str(HospitalListNewOrderPage);#string type here
    print(x);
    hospitalLabel = x.split(",",1);
    hospitalNameTemp=hospitalLabel[1];
    #print(hospitalNameTemp);
    hospitalIDTemp=hospitalLabel[0];
    #print(hospitalIDTemp);
    HospitalID=hospitalIDTemp.split("(");
    hospitalId=int(HospitalID[1]);

    HospitalLabelForEditPage.pack();
    HospitalLabelForEditPage.place(x=20,y=80);

    FindActivityByHospitalID=f'SELECT Fixed_activity_level FROM hospital where idhospital="{hospitalId}"';
    cursor.execute(FindActivityByHospitalID);
    ActivitiyBySpecificOrder = cursor.fetchall();
    TempID=[i[0] for i in ActivitiyBySpecificOrder];#find index number in a list of tuple
    ActivitiyBySpecificOrder2=TempID[0];
    print("Activity of Hospital selected:",ActivitiyBySpecificOrder2);


    #Create tree/table for the Edit page
    EditTree = ttk.Treeview(EditPage,height=15,selectmode="browse");#select=browse means can choose only 1 record at the time
    EditTree['columns']= ("ID","Amount","Injection time")
    EditTree.pack();
    EditTree.place(x=170,y=130);

        #Edit page/tree scrollbar -vertical
    EditOrderPageScroll = Scrollbar(EditPage, orient="vertical", command=EditTree.yview);
    EditOrderPageScroll.place(x=515, y=130, height=330)
    EditTree.configure(yscrollcommand = EditOrderPageScroll.set);


#Foramte Columns
    EditTree.column("#0",width=0,minwidth=0);
    EditTree.column("ID",anchor=W,width=80,minwidth=25);
    EditTree.column("Amount",anchor=CENTER,width=140,minwidth=25);
    EditTree.column("Injection time",anchor=W,width=120,minwidth=25);

    #Define headers/titles in table
    EditTree.heading("#0", text="Label",anchor=W);
    EditTree.heading("ID", text="Dose number",anchor=W);
    EditTree.heading("Amount", text=f"Amount of Doses ({ActivitiyBySpecificOrder2})",anchor=CENTER);
    EditTree.heading("Injection time", text="Injection time",anchor=W);

    #get order detail from DB by hospitalID and Date
    cursor = db.cursor();
    FindActivityByHospitalID=f'SELECT Fixed_activity_level FROM hospital where idhospital="{hospitalId}"';
    cursor.execute(FindActivityByHospitalID);
    ActivitiyBySpecificOrder = cursor.fetchall();
    TempID=[i[0] for i in ActivitiyBySpecificOrder];#find index number in a list of tuple
    ActivitiyBySpecificOrder2=TempID[0];
    print("Activity of Hospital selected:",ActivitiyBySpecificOrder2);


    SearchSpecOrderQueryByDoubleclick=f'SELECT DoseNumber,amount,Injection_time FROM orders where hospitalID="{hospitalId}" AND Date="{DateSelected}"';
    cursor.execute(SearchSpecOrderQueryByDoubleclick);
    OrderDatatoSpecificOrder = cursor.fetchall();
    print(f"order selected: {OrderDatatoSpecificOrder}");
    ListOfInjectionTime=[];
    ListOfAmount=[];
    IdOrderPK=[];
    #output orders main data from DB to the orders tree
    for record in OrderDatatoSpecificOrder:
        EditTree.insert(parent='', index='end',values=(record[0],record[1],record[2]));#record[0]=Id,record[1]=amount,record[2]=injection time
        ListOfInjectionTime.append(record[2]);
        ListOfAmount.append(record[1]);
        IdOrderPK.append((record[0]));
    #EditTree.pack();
    print(ListOfInjectionTime);
    print(ListOfAmount);
    # declaring string variable for storing amount
    amountVar=AmountOfDosesSelected;
    # declaring string variable for storing time interval
    TimeIntervals="30";
    # declaring string variable for storing time
    # HoursVar = "1";
    # MinutesVar = "1";
    global OrderID,DosesSelectedEvent,idCounter;
    #global DosesSelectedEvent;
    OrderID=0;
    #global idCounter;
    ListofVal=["","","","","","",""];
    #global hospitalId;
    #global OrderID;


    #Get Time varibles avent,hous and minutes
    Time_Intervals=TimeIntervals;
    # Minutes_Var=MinutesVar;
    # Hours_Var=HoursVar;
    #get amount event variable
    #message box if not try to click next if inputs are empty
    amount=amountVar;
    IntAmount=(int(amount));
    ListofVal[0]=idCounter=1;
    ListofVal[1]=(IntAmount / IntAmount);
    # ListofVal[2]=int(Hours_Var);
    # ListofVal[3]=int(Minutes_Var);
    ListofVal[4]=IntAmount;
    ListofVal[5]=hospitalId;
    ListofVal[6]=int(Time_Intervals);

    OrderID+=1;#counterID=counterID+1
    #print(OrderID);


    #Update Date in GUI
    ChoosenDateForManaulOrder=DateSelected;
    print(ChoosenDateForManaulOrder);
    #copy and past date event to page number 2
    dateLabel2=Label(EditPage, text= ChoosenDateForManaulOrder, bg="white", font=('Helvetica 14'));
    dateLabel2.pack();
    dateLabel2.place(x=20,y=110);


    #
    # #Create Time range input/Entry
    # TimerangeLabel = Label(NewOrderMainPage, text="Time Range/Intervals",bg="white");
    # TimerangeLabel.place(x=400, y=200);
    # TimerangeLabelEntry = Entry(NewOrderMainPage,textvariable=TimeIntervals,font=("Halvetica",12));
    # TimerangeLabelEntry.config(width=7);#width of window
    # TimerangeLabelEntry.insert(0, '');
    # TimerangeLabelEntry.pack();
    # TimerangeLabelEntry.place(x=400, y=240);

    #Title of page
    NewOrdersecondaryLabel=Label(EditPage, text="Edit order",bg="#F0F3F4", font=('Helvetica 18 bold'), fg='#034672');
    NewOrdersecondaryLabel.pack();
    NewOrdersecondaryLabel.place(x=270,y=27);




    # #Create a Cancel button
    # global CancelImg;
    # CancelIcon = Image.open("./Images/CancelButton.png");
    # resized_Cancel_Icon = CancelIcon.resize((100,50), Image.ANTIALIAS);
    # CancelImg = ImageTk.PhotoImage(resized_Cancel_Icon);
    CancelButton2=Button(EditPage,text="Done editing",command=lambda: [EditPage.destroy()]);#close window-not working
    CancelButton2.pack();
    CancelButton2.place(x=390, y=520);

#Update values,amount and injection time####################################################
    global ClickingRowsCounter;
    ClickingRowsCounter=0;
    TempListForUpdateOrderValues=["",""];#List for saving

    #############################Amount quantity Event ######################

    AmountList = [
        "0","1","2","3","4","5"
    ]

    status = tk.StringVar()
    status.set("0")

    #Catch event
    def treeAmountSelect(event):
        global ClickingRowsCounter;

        row = EditTree.focus();#catch row item

        #Count numbers of chnages in order editing
        ClickingRowsCounter=ClickingRowsCounter+1;
        if ClickingRowsCounter>1:
            messagebox.showinfo(title="Info message", message="You cant edit more then 1 record at the time");

        #Enble  Injection-time and amount dropdown menues
        dropDownInjectionT_M.configure(state="normal ");
        dropDownAmountM.configure(state="normal ");

        dataofchoosnenRowListEditTree=row;
        print(dataofchoosnenRowListEditTree);
        DataOfRowSelectedDicEditTree=EditTree.item(dataofchoosnenRowListEditTree);
        DataOfRowSelectedList=DataOfRowSelectedDicEditTree['values'];
        print(DataOfRowSelectedList);
        IidSelected=DataOfRowSelectedList[0];
        AmountSelected=DataOfRowSelectedList[1];
        InjectionTimeSelected=DataOfRowSelectedList[2];

    #print(row)
        if row:
            status.set(EditTree.set(row, 'Amount'));

    EditTree.bind('<<TreeviewSelect>>', treeAmountSelect);

    def set_amount(AmountValue):
        #global DosesSelectedEvent;
        print(AmountValue);
        TempListForUpdateOrderValues[1]=int(AmountValue);
        #print(TempListForUpdateOrderValues[1]);
        #DosesSelectedEvent=AmountValue;
        row = EditTree.focus();

        if row:
            #try:
            #     cursor = db.cursor(buffered=True);
            #     UpdateSQlQuery=f"UPDATE  orders SET injection_time='{InjectionTimeSelected}',amount='{AmountValue}',batchID='{0}',DecayCorrected='{0}'  WHERE idorders = '{IidSelected}';";
            #     cursor.execute(UpdateSQlQuery);
            #     print("DB updated successfully ");
            #     db.commit();
            #     cursor.close();
            # except Exception as e:
            #     logging.error(traceback.format_exc())
            #     print("Error-Order was not updated-please check MySQL")
            EditTree.set(row, '1', AmountValue);


    dropDownAmountM = ttk.OptionMenu(EditPage, status, "0", *AmountList, command=set_amount);
    dropDownAmountM.pack();
    dropDownAmountM.place(x=200, y=490);
    #Change Amount time manual label
    ChangeAmountLabel=Label(EditPage, text="Change Amount : ", font=('Helvetica 12'));
    ChangeAmountLabel.pack();
    ChangeAmountLabel.place(x=20,y=490);




    # ############################Injection Time event#########################
    TimeList = [
        "06:00","06:30","07:00","07:30","08:00","08:30","09:00","09:30","10:00","10:30","11:00",
        "11:30","12:00","12:30","13:00","13:30","14:00","14:30","15:00","15:30","16:00","16:30",
        "17:00","17:30","18:00","18:30","19:00","19:30","20:00"
    ]

    Timeselected = tk.StringVar()
    Timeselected.set("00:00")
    #print(status);
    #Catch Injection  event
    def InjectionTimeselect(event):
        global DoseNum_Selected,AmountSelected,InjectionTimeSelected,DosesSelectedEvent;

        global ClickingRowsCounter;

        row = EditTree.focus();
        ClickingRowsCounter=ClickingRowsCounter+1;
        if ClickingRowsCounter>1:
            messagebox.showinfo(title="Info message", message="You cant edit more then 1 record at the time");


        #Enble  Injection-time and amount dropdown menues
        dropDownInjectionT_M.configure(state="normal ");
        dropDownAmountM.configure(state="normal ");

        dataofchoosnenRowListEditTree=row;
        #print("Row selected",dataofchoosnenRowListEditTree);
        DataOfRowSelectedDicEditTree=EditTree.item(dataofchoosnenRowListEditTree);
        DataOfRowSelectedList=DataOfRowSelectedDicEditTree['values'];
        print("Row/Values selected",DataOfRowSelectedList);
        DoseNum_Selected=DataOfRowSelectedList[0];
        AmountSelected=DataOfRowSelectedList[1];
        TempListForUpdateOrderValues[1]=AmountSelected;#insert current amount selected to the list-Defualt
        InjectionTimeSelected=DataOfRowSelectedList[2];
        TempListForUpdateOrderValues[0]=InjectionTimeSelected;#insert current time_injection selected to the list-Defualt

        if row:
            Timeselected.set(EditTree.set(row, 'Injection time'));

    EditTree.bind('<<TreeviewSelect>>', InjectionTimeselect)

    def setInjectionTime(CurValueForTime):
        global DoseNum_Selected,AmountSelected,InjectionTimeSelected,DosesSelectedEvent;
        print(CurValueForTime);
        TempListForUpdateOrderValues[0]=CurValueForTime;
        #print(f'Injection time selected:{TempListForUpdateOrderValues[0]}');
        row = EditTree.focus();
        # dataofchoosnenRowListEditTree=row;
        # #print(dataofchoosnenRowListEditTree);
        # DataOfRowSelectedDicEditTree=EditTree.item(dataofchoosnenRowListEditTree);
        # DataOfRowSelectedList=DataOfRowSelectedDicEditTree['values'];
        # print(DataOfRowSelectedList);
        if row:
            EditTree.set(row, '2', CurValueForTime)

    dropDownInjectionT_M = ttk.OptionMenu(EditPage, Timeselected, "00:00", *TimeList, command=setInjectionTime);
    dropDownInjectionT_M.pack();
    dropDownInjectionT_M.place(x=200, y=460);
    #Change injecion time manual label
    ChangeTImeIjectionLabel=Label(EditPage, text="Change Time-Injection : ", font=('Helvetica 12'));
    ChangeTImeIjectionLabel.pack();
    ChangeTImeIjectionLabel.place(x=20,y=460);


    def enterToDB_UpdateOrder():
        """Function for submitting/save changes from edit order page"""
        global DoseNum_Selected,ClickingRowsCounter;
        ClickingRowsCounter=0;
        #Disable  Injection-time and amount dropdown menues
        dropDownInjectionT_M.configure(state="disabled")
        dropDownAmountM.configure(state="disabled")
        #print(TempListForUpdateOrderValues[0]);

        try:
            cursor = db.cursor(buffered=True);
            UpdateSQlQuery=f"UPDATE  orders SET injection_time='{TempListForUpdateOrderValues[0]}',amount='{TempListForUpdateOrderValues[1]}' WHERE hospitalID = '{hospitalId}' and Date = '{DateSelected}' and DoseNumber={DoseNum_Selected};";
            cursor.execute(UpdateSQlQuery);
            print("DB updated successfully ");
            db.commit();
            cursor.close();
        except Exception as e:
            logging.error(traceback.format_exc())
            messagebox.showerror("Error message","Error !");
            print("Error-Order was not updated-please check MySQL")


        #Function to insert data into My-SQL Db
        #EditPage.destroy();#Close import file-manual window
        updateOrdersTreeMainPageOutputOnly();#Refresh/Update Main page
        #OrdersTree.pack();         #open order main page immedaitly
#

    #Create ADD button
    # global addImg;
    # AddFileIcon = Image.open("./Images/AddButton.png");
    # resized_add_Icon = AddFileIcon.resize((100,50), Image.ANTIALIAS);
    # addImg = ImageTk.PhotoImage(resized_add_Icon);
#
    AddButton1=Button(EditPage,text="Update record",command=enterToDB_UpdateOrder);
    AddButton1.pack();
    AddButton1.place(x=180, y=520);

    #Disable  Injection-time and amount dropdown menues
    dropDownInjectionT_M.configure(state="normal")
    dropDownAmountM.configure(state="normal")




    #Create ADD row button+icon


    # defining a function that will
    # print them on the screen
    #rowTree = StringVar();
    def addRowFunc():
        #global idCounter;
        rowTreetoAdd=(ListofVal[0],ListofVal[1],ListofVal[2]);
        EditTree.insert("", "end", values=rowTreetoAdd);
        ListofVal[0]=ListofVal[0]+1;
        ListofVal[4]+=1;#current amount= courrent amount+1

    def removeRawFunc():
        global DoseNum_Selected;
        #rowTree=rowTree.get();
        #for i,j in zip(range(IntAmount),range(BeginigHour,IntAmount)):
        rawSelectedToDelete=EditTree.selection();
        for rawselected in rawSelectedToDelete:
            try:
                EditTree.delete(rawselected);
                cursor = db.cursor(buffered=True);
                DeleteRecordOrdersTable_query = f'DELETE FROM orders WHERE idorders = {DoseNum_Selected}';
                cursor.execute(DeleteRecordOrdersTable_query);
                print("Remove record from DB sucssful ");
                db.commit();
                cursor.close();
            except Exception as e:
                logging.error(traceback.format_exc())
                print("Error with,record not delete, check MySQL");

        ListofVal[4]=ListofVal[4]-1;#current amount= current amount-1
    #amountVar.set("");




    ####################Buttons for new order-manual page##########################
    # Remove button (Icon) - List
    global imgDelete2;
    deleteIcon2 = Image.open("./‏‏deleteIcon.png");
    resizedDeleteIcon2 = deleteIcon2.resize((25,25), Image.ANTIALIAS);
    imgDelete2 = ImageTk.PhotoImage(resizedDeleteIcon2);
    deleteButton2=Button(EditPage, image=imgDelete2, borderwidth=0,command=removeRawFunc);
    deleteButton2.pack();
    deleteButton2.place(x=460, y=98);


    #Creaet Label Add row
    global addROWImg;
    AddrowLabel=Label(EditPage, text="Add row",bg="white", font=('Helvetica 14'));
    AddrowLabel.pack();
    AddrowLabel.place(x=280,y=97);
    #Add row image+button
    AddrowIcon = Image.open("./addIcon.png");
    resized_add_Row = AddrowIcon.resize((25,25), Image.ANTIALIAS);
    addROWImg = ImageTk.PhotoImage(resized_add_Row);
    AddRowButton=Button(EditPage,image=addROWImg, borderwidth=0,command=addRowFunc);
    AddRowButton.pack();
    AddRowButton.place(x=250, y=98);

#Double click on  main order page tree event
OrdersTree.bind('<<TreeviewOpen>>', UpdateOrder)

###############################End of update page#################################################

###########################Spacial buttons#################################################

#Create a button for import orders files (Excel or Word)
# ImportFileIcon = Image.open("ImportFile2.png")
# ImportFileIcon = Image.open("ImportFile2.png")
# resized_Edit_Icon = ImportFileIcon.resize((80,20), Image.ANTIALIAS)
# img_Edit = ImageTk.PhotoImage(resized_Edit_Icon)
importFileButton=Button(ordersFrame, text="Import file",command=importFileFunc)
importFileButton.pack()
importFileButton.place(x=250, y=65)



#Create New order button
# NewOrderIcon = Image.open("./Images/AddnewOrder2.png")
# resizedNewOrderIconIcon = NewOrderIcon.resize((120,20), Image.ANTIALIAS)
# NewOrderIconimg = ImageTk.PhotoImage(resizedNewOrderIconIcon)
editButton=Button(ordersFrame, text="Add new order",command=PopUpForNewOrder);
editButton.pack();
editButton.place(x=100, y=65);




root.mainloop();